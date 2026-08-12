from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path("/Users/jk0307/Documents/GitHub/iclo/iclo-us-employee-dashboard")
OUT = ROOT / "output/proposal-v4"
ICLO_LOGO = ROOT / "output/booth/ICLO-Logo-Color-Transparent.png"
SNOW_LOGO = ROOT / "output/booth/assets/Snowflake-Corporate-Logo-Blue-Transparent-v3.png"

NAVY = "1D3155"
TEAL = "008A94"
SKY = "29B5E8"
CORAL = "FF6F73"
AMBER = "E09A19"
GRAY = "63738A"
LIGHT = "F3F7FA"
PALE_BLUE = "EAF7FC"
PALE_TEAL = "EAF6F5"
PALE_CORAL = "FFF0F0"
PALE_AMBER = "FFF7E8"
WHITE = "FFFFFF"


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def set_run(run, size=None, bold=None, color=None, font="Arial"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document, internal: bool):
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(9.3)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [("Title", 28, NAVY), ("Heading 1", 21, NAVY), ("Heading 2", 13, TEAL), ("Heading 3", 10.5, NAVY)]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header
    t = header.add_table(1, 3, width=Mm(174))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Mm(28)
    t.columns[1].width = Mm(118)
    t.columns[2].width = Mm(28)
    t.cell(0, 0).paragraphs[0].add_run().add_picture(str(ICLO_LOGO), width=Mm(24))
    p = t.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ICLO × SNOWFLAKE  |  " + ("INTERNAL REVIEW" if internal else "JOINT VALIDATION PROPOSAL"))
    set_run(r, 7.5, True, GRAY)
    t.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    t.cell(0, 2).paragraphs[0].add_run().add_picture(str(SNOW_LOGO), width=Mm(24))
    for cell in t.rows[0].cells:
        set_cell_margins(cell, 0, 0, 40, 0)

    footer = section.footer
    footer.paragraphs[0].text = ""
    ft = footer.add_table(1, 2, width=Mm(174))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    ft.autofit = False
    ft.columns[0].width = Mm(140)
    ft.columns[1].width = Mm(34)
    p = ft.cell(0, 0).paragraphs[0]
    r = p.add_run("ICLO  |  August 2026")
    set_run(r, 7, False, GRAY)
    p2 = ft.cell(0, 1).paragraphs[0]
    add_page_number(p2)
    for cell in ft.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)


def page_break(doc: Document):
    doc.add_page_break()


def eyebrow(doc: Document, text: str, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run(r, 8.5, True, color)


def title(doc: Document, text: str, subtitle: str | None = None):
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(12)
        r = p2.add_run(subtitle)
        set_run(r, 11.5, False, GRAY)


def section_title(doc: Document, eyebrow_text: str, heading: str, takeaway: str):
    eyebrow(doc, eyebrow_text)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(5)
    p.add_run(heading)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(takeaway)
    set_run(r, 10.5, False, GRAY)


def paragraph(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, 9.3, True, NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_run(r, 9.3, False, NAVY)
    else:
        r = p.add_run(text)
        set_run(r, 9.3, False, NAVY)
    return p


def bullets(doc: Document, items: Iterable[str], color=NAVY):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-2.5)
        r = p.add_run(item)
        set_run(r, 9.1, False, color)


def callout(doc: Document, label: str, text: str, fill: str, accent: str):
    table = doc.add_table(1, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(43)
    table.columns[1].width = Mm(129)
    c0, c1 = table.rows[0].cells
    shade(c0, accent)
    shade(c1, fill)
    for c in (c0, c1):
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c, 130, 150, 130, 150)
    p = c0.paragraphs[0]
    r = p.add_run(label.upper())
    set_run(r, 7.8, True, WHITE)
    p = c1.paragraphs[0]
    r = p.add_run(text)
    set_run(r, 8.7, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def process_row(doc: Document, stages: Sequence[tuple[str, str]], colors: Sequence[str] | None = None):
    colors = colors or [PALE_CORAL, PALE_BLUE, PALE_AMBER, PALE_TEAL]
    table = doc.add_table(1, len(stages))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (head, body) in enumerate(stages):
        cell = table.cell(0, idx)
        shade(cell, colors[idx % len(colors)])
        set_cell_margins(cell, 180, 160, 180, 160)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(head + "\n")
        set_run(r, 9.4, True, NAVY)
        r = p.add_run(body)
        set_run(r, 8.2, False, GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def data_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None, font_size=7.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for col, width in zip(table.columns, widths):
            col.width = Mm(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        shade(cell, TEAL)
        set_cell_margins(cell, 120, 100, 120, 100)
        r = cell.paragraphs[0].add_run(text)
        set_run(r, font_size, True, WHITE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            shade(cell, WHITE if row_idx % 2 == 0 else LIGHT)
            set_cell_margins(cell, 100, 100, 100, 100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cell.paragraphs[0].add_run(text)
            set_run(r, font_size, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def cover(doc: Document, internal: bool):
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    table = doc.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    set_cell_margins(cell, 500, 450, 500, 450)
    p = cell.paragraphs[0]
    r = p.add_run("ICLO × SNOWFLAKE\n")
    set_run(r, 12, True, SKY)
    p.add_run("\n")
    heading = "임직원 치아보험 이용과\n청구로 확인하는 근거 레이어" if internal else "Employee Dental-Benefit Navigation\nand a Claims-Verified Evidence Layer"
    r = p.add_run(heading + "\n")
    set_run(r, 25, True, WHITE)
    p.add_run("\n")
    sub = "내부 검토용 제안서 | 미국 치아보험·금융 용어 설명 포함" if internal else "Snowflake HLS GTM & Architecture Joint-Validation Proposal"
    r = p.add_run(sub + "\n")
    set_run(r, 11, False, WHITE)
    p.add_run("\n")
    r = p.add_run("August 2026")
    set_run(r, 8.5, False, SKY)
    doc.add_paragraph()
    callout(
        doc,
        "읽는 방법" if internal else "POSITIONING",
        "B-prime의 미국 치아보험 문제 제기로 시작하고, A안의 Snowflake 협업 제안을 본문으로 사용합니다." if internal else "U.S. dental-insurance friction opens the story; the Snowflake governed-evidence collaboration is the main proposal.",
        PALE_BLUE,
        SKY,
    )


def build_english():
    doc = Document()
    configure(doc, internal=False)
    cover(doc, internal=False)

    page_break(doc)
    section_title(doc, "Executive decision brief", "What This Proposal Is Designed to Achieve", "The ask is a validation path, not a product endorsement or generic customer introduction.")
    process_row(doc, [
        ("1. Business fit", "Agree on the most natural Snowflake HLS sales play"),
        ("2. Architecture", "Secure an HLS technical working session"),
        ("3. Account evidence", "Validate one payer/TPA account team before introductions"),
    ], [PALE_TEAL, PALE_BLUE, PALE_AMBER])
    paragraph(doc, "ICLO helps employees understand and use their dental benefits, navigate to appropriate in-network care, and converts the journey into claims-verified, aggregate evidence for the employer.")
    paragraph(doc, "Short form: Employee dental-benefit navigation in the front; a governed evidence layer in the back.", "Short form:")
    callout(doc, "CONFIRMED ICLO DIRECTION", "Lead with employee dental-benefit navigation and employer aggregate evidence - not oral-image AI.", PALE_BLUE, NAVY)
    callout(doc, "PRODUCT DESIGN PRINCIPLE", "Keep model-derived signals separate from automated clinical routing; keep employee-level health signals out of employer views.", PALE_TEAL, TEAL)
    callout(doc, "HYPOTHESIS / TO VALIDATE", "The pattern can become repeatable across payer/TPA accounts that hold usable dental eligibility and claim-line data.", PALE_AMBER, AMBER)
    callout(doc, "REQUEST TO SNOWFLAKE", "Provide sales-play fit, named internal stakeholders, architecture decision support, account-team validation and written readiness criteria.", PALE_BLUE, SKY)

    page_break(doc)
    section_title(doc, "Business problem", "Dental Insurance Does Not Automatically Become Usable Care", "Eligibility is only the first step; network, plan rules and real provider availability shape access and employee cost.")
    process_row(doc, [
        ("Dental insurance", "Employee or dependent is eligible"),
        ("Network", "Provider choice and economics vary"),
        ("Dental plan", "Deductible, coinsurance, annual maximum"),
        ("Availability", "Directory status and actual booking"),
        ("Usable care", "Expected out-of-pocket cost is understood"),
    ])
    bullets(doc, [
        "Provider choice may be constrained or economically penalized by network and plan rules.",
        "A provider directory may not reflect current appointment availability.",
        "Pre-treatment estimates can differ from the final explanation of benefits and claim result.",
        "An employer-sponsored relationship can create a contractual path to plan-specific data, subject to the employer's data rights and TPA/carrier cooperation.",
    ])
    callout(doc, "ONE EMPLOYEE'S JOURNEY", "A covered employee may still need to identify the right network, understand cost-sharing, find an available dentist and wait for the claim before the actual benefit use is clear.", PALE_BLUE, SKY)
    callout(doc, "BOUNDARY", "Employer contracting does not automatically provide complete provider, eligibility or claims data.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "Product", "ICLO Connects Dental-Benefit Context to Employee-Initiated Action", "ICLO is an employee benefits workflow with a claims-verified evidence layer - not another teledentistry consultation app.")
    data_table(doc, ["Layer", "What it contains", "Operating discipline"], [
        ["Employee experience", "Self-reported concern, benefit understanding, optional oral image, provider search, requested support", "Employee action or human-assisted navigation remains explicit"],
        ["Benefit and provider context", "Eligibility, network, plan rules, deductible, coinsurance, annual maximum, available appointments", "Only where contractually obtainable"],
        ["Claims evidence", "Member-month denominator, preventive completion, allowed, plan-paid, employee OOP, treatment mix, run-out", "Employer output is aggregate only"],
        ["Shadow signal", "URI, metadata, model version and approved derived signal", "No raw disease probabilities; no automated provider, urgency or treatment-pathway decision"],
    ], [34, 82, 56], 7.5)
    paragraph(doc, "ICLO is not: a direct-to-consumer diagnostic AI, an underwriting or employee-risk scoring product, a provider-referral marketplace, or a first-year claims-savings promise.", "ICLO is not:")
    callout(doc, "PROVISIONAL REDLINE", "Separate signal, navigation and referral-event tables; timestamp employee action and consent; maintain purpose-based access and model-version lineage, pending U.S. regulatory counsel review.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Architecture", "Snowflake as the Governed Evidence and Collaboration Plane", "Structured eligibility, plan, event and dental-claims data converge in Snowflake; raw oral images follow a separate controlled path.")
    paragraph(doc, "Structured evidence path", "Structured evidence path")
    process_row(doc, [
        ("App events", "Consent and employee actions"),
        ("HRIS eligibility", "Employee and dependent denominator"),
        ("Plan + provider context", "Network and benefit rules"),
        ("TPA/carrier claims", "Dental claim lines and run-out"),
        ("Snowflake evidence layer", "Canonical model, controls, aggregate outcomes"),
    ], [PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_TEAL])
    paragraph(doc, "Raw-image path", "Raw-image path")
    process_row(doc, [
        ("Raw oral image", "Controlled U.S. object storage / PHI vault"),
        ("Shadow inference", "Model endpoint and validation outside this proposal"),
        ("Snowflake metadata", "URI, metadata, model version, approved signal only"),
        ("Employer dashboard", "Aggregate views; no employee-level routing or health signal"),
    ], [PALE_CORAL, PALE_AMBER, PALE_BLUE, PALE_TEAL])
    callout(doc, "PLATFORM PREREQUISITE", "If PHI is processed in Snowflake, Business Critical and an executed BAA are platform prerequisites. They do not make ICLO as a whole HIPAA-compliant.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Economics", "The Economics Begin with a J-Curve - Not a Year-One Savings Promise", "Preventive utilization and newly discovered treatment can increase dental plan-paid claims before treatment-mix change becomes measurable.")
    data_table(doc, ["Period / segment", "What may happen", "How to interpret it"], [
        ["Year 1", "Preventive visits rise; untreated conditions are discovered", "Allowed, plan-paid and employee OOP may move differently"],
        ["Years 2-3", "Treatment mix and claims run-out become more observable", "Direction remains a customer-specific hypothesis"],
        ["Low turnover", "More time for the employer to observe multi-year effects", "A multi-year value case may be more relevant"],
        ["High turnover", "Untreated need may re-enter with new hires", "Employee experience, recruitment and retention may be the stronger case"],
    ], [32, 76, 64], 7.7)
    paragraph(doc, "Year-one utilization and plan-paid claims may rise before any longer-term change in treatment mix becomes measurable. ICLO will not contract on a year-one savings guarantee.")
    paragraph(doc, "Planning inputs include population age, geography, network overlap, provider availability, turnover, plan design, historical claims, paid/OOP allocation, untreated-need assumptions and claims run-out. Outputs include Year 1-3 scenarios, turnover sensitivity, network-access sensitivity and explicit uncertainty.")
    callout(doc, "HYPOTHESIS", "Simulation values are planning heuristics to be calibrated with customer data - not actual ICLO outcomes.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Responsibilities", "What Snowflake Can Support - and What Remains Outside Its Scope", "The platform can enable governed evidence and collaboration; it does not decide ICLO's intended use, data rights or legal sufficiency.")
    data_table(doc, ["Snowflake can support", "ICLO / customer responsibility", "Outside the Snowflake ask"], [
        ["Governed storage and compute; canonical model; reconciliation", "Source mapping, product logic and operational controls", "Employer/TPA data-rights determination"],
        ["Role-based access, masking, row policies and access history", "Consent UX, purpose metadata and incident operations", "Legal sufficiency of employee consent"],
        ["Aggregate outcome views and claims-confirmed calculations", "Program attribution logic and customer reporting", "First-year savings validation"],
        ["Same-region share where available; controlled staging otherwise", "Connector and source-system implementation", "Cross-region legal conclusions"],
        ["Workload sizing and post-credit unit economics", "Commercial assumptions and usage monitoring", "FDA status, licensure or referral-law analysis"],
    ], [58, 58, 56], 7.5)
    paragraph(doc, "Native App, multi-party Clean Room and federated learning are conditional options. They are not prerequisites for the first employer pilot.")

    page_break(doc)
    section_title(doc, "Evidence design", "The Dashboard Makes Data Quality and Privacy Visible", "The current dashboard is synthetic visual proof, not a performance claim.")
    data_table(doc, ["Visible control", "Why the buyer needs it"], [
        ["Synthetic data - illustrative only", "Prevents the demo from being mistaken for an ICLO outcome"],
        ["Employees lens / members lens", "Separates the application-user denominator from covered member-month cost denominator"],
        ["Eligibility and claims dates", "Shows which plan period and source cutoff each result uses"],
        ["Claims lag and completeness", "Makes run-out and missing-data uncertainty visible"],
        ["Aggregate only; n >= 20 suppression", "Reduces re-identification risk and keeps individual PHI out of employer views"],
    ], [56, 116], 8.0)
    callout(doc, "MESSAGE", "This is not a performance claim. It demonstrates how denominator, freshness, privacy and claims completeness become visible to the buyer.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "Account ecosystem", "U.S. Dental Benefits Are a Multi-Party, Account-Specific Workflow", "The employer sponsors the benefit; a dental carrier or TPA administers it; providers deliver care; claims confirm what happened.")
    process_row(doc, [
        ("Employer / benefits team", "Sponsors the dental benefit and selects funding + administrator"),
        ("Dental carrier / TPA", "Runs network, benefit rules, eligibility and/or claims"),
        ("Employee / dental provider", "Uses care; the provider submits the claim"),
        ("Claims + evidence", "Confirms visits, procedures, allowed, plan-paid and OOP"),
    ])
    data_table(doc, ["Funding model", "Who bears claims risk", "Who may administer", "Why it matters to ICLO"], [
        ["Fully insured", "Dental carrier / insurer", "Usually the carrier", "Plan, network and claims data follow the carrier and legal entity."],
        ["Self-funded / ASO", "Employer", "Carrier or independent TPA", "The administrator may hold eligibility and claim lines; employer data rights still require validation."],
    ], [34, 38, 42, 58], 7.2)
    paragraph(doc, "Large brands such as Delta Dental and Blue plans illustrate why the exact legal entity matters: regional or member companies can have different procurement, systems and vendors. These are examples of the rule, not target-account assumptions.")
    callout(doc, "ACCOUNT-BY-ACCOUNT DISCIPLINE", "Start with funding model, administrator, legal entity, cloud/region, dental workload and current vendor. Do not claim that 'many TPAs are on Snowflake' without account evidence.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "Partnership request", "What We Need from Snowflake GTM and Architecture", "The GTM specialist is an industry-fit, internal-coordination and account-path sponsor - not the implementation owner.")
    data_table(doc, ["Priority", "Requested output"], [
        ["Immediate", "Named HLS architect and solution engineer; U.S. account/region/Business Critical/BAA path; written credits and support scope; 2.5K/10K/25K sizing; security reference architecture"],
        ["Next", "Payer/carrier-ASO/TPA named-account map; one account-team use-case validation; Direct Share versus SFTP/API decision pattern; partner checklist; co-sell milestones"],
        ["Later / conditional", "Native App feasibility; multi-party Clean Room; federated-learning feasibility; joint case study; optional Snowflake employee-benefits design partnership, subject to fit"],
    ], [38, 134], 7.8)
    bullets(doc, [
        "Which Snowflake HLS sales play fits most naturally?",
        "Is the stronger Snowflake story the employee application or the governed payer/TPA evidence layer behind it?",
        "What must ICLO prove before an account team would actively introduce or co-sell?",
        "Should the path be Snowflake for Startups, Partner Network, Powered by Snowflake or an ISV/Native App route?",
    ])

    page_break(doc)
    section_title(doc, "Execution", "A 90-Day Joint Validation Plan", "The desired result is a validated architecture and one account-team review - not a co-sell commitment.")
    process_row(doc, [
        ("Days 0-30 | Align", "Sales play, named stakeholders, U.S. account/region/BAA path, architecture workshop"),
        ("Days 31-60 | Design", "Ingestion pattern, population sizing, security reference, account matrix v1"),
        ("Days 61-90 | Validate", "One payer/TPA account-team review, partner checklist, next/no-go decision"),
        ("Conditional | Expand", "Native App, Clean Room, federated learning, joint story or design partner"),
    ])
    paragraph(doc, "We are not asking Snowflake to validate our dental AI or make ICLO compliant. We are asking Snowflake to help determine whether this can become a repeatable, governed payer/TPA data-collaboration pattern - and what ICLO must prove to make the partnership commercially real.")

    page_break(doc)
    section_title(doc, "Meeting kit", "90-Second Opening Statement", "An external verbal opening for a Snowflake HLS GTM specialist.")
    script = (
        "ICLO is building an employee dental-benefit navigation and claims-verified evidence layer. In the front, we help employees understand their dental benefits, find appropriate in-network care and request appointment support. In the back, we connect eligibility, plan context, app events and dental claims so the employer can see aggregate, privacy-governed evidence - without seeing individual health signals.\n\n"
        "In a U.S. employer-sponsored dental benefit, the employer sponsors the program, while a dental carrier or TPA may administer eligibility, network rules and claims. Employees use care, providers submit claims, and those claims confirm what happened. The problem is not coverage alone: network rules, deductibles, annual maximums, provider availability and claims lag all shape whether a covered employee can use the benefit and understand the cost. We also do not promise first-year claims savings; preventive use and newly discovered treatment can raise plan-paid claims before longer-term change becomes measurable.\n\n"
        "We see Snowflake as the governed evidence and collaboration plane behind this workflow. We are not asking Snowflake to validate our dental AI, make ICLO compliant or introduce us broadly to customers. We are asking for three things: agreement on the most natural HLS sales play, a technical working session on the payer/TPA data pattern, and an account-by-account path to validate whether dental eligibility or claim-line data actually exists in Snowflake. Our goal is to learn what ICLO must prove for this to become repeatable and commercially real."
    )
    paragraph(doc, script)
    paragraph(doc, "45-minute agenda", "45-minute agenda")
    data_table(doc, ["Minutes", "Topic", "Decision / output"], [
        ["0-8", "Business fit", "Primary HLS sales play and strongest Snowflake story"],
        ["8-16", "Account ecosystem", "Account types and evidence required for a named-account map"],
        ["16-29", "Architecture", "Share/ingestion, tenancy, PHI boundary, workload-sizing questions"],
        ["29-38", "Partnership path", "Startup/partner/ISV path and readiness criteria"],
        ["38-45", "Next actions", "Named owners, dates and required written outputs"],
    ], [20, 48, 104], 7.7)

    page_break(doc)
    section_title(doc, "Action register", "Owners and Required Outputs", "Every next step should end with a named owner, due date and artifact.")
    data_table(doc, ["Action", "Snowflake owner", "ICLO owner", "Due date", "Required output"], [
        ["Confirm HLS sales-play fit", "HLS GTM specialist", "CEO / GTM lead", "Meeting + 3 business days", "One written primary/secondary sales-play statement"],
        ["Name technical working-session team", "HLS Industry Architect", "Product / data lead", "Meeting + 5 business days", "Architect, SE and security/compliance contacts"],
        ["Confirm U.S. account, region, Business Critical and BAA path", "Account / security owner", "Security lead", "Meeting + 10 business days", "Written decision path and prerequisites"],
        ["Run 2.5K / 10K / 25K workload sizing", "Solution Engineer", "Data lead", "Day 30", "Assumptions, credits, post-credit unit economics"],
        ["Build named-account matrix v1", "HLS GTM + account teams", "GTM lead", "Day 45", "Legal entity, dental workload, region, owner and readiness"],
        ["Validate one payer/TPA account-team use case", "Named account owner", "CEO / product lead", "Day 75", "Written architecture and unmet-need feedback"],
        ["Decide partner path", "Startup / Partner lead", "CEO / partnerships", "Day 90", "Checklist; proceed/defer decision"],
    ], [48, 32, 30, 28, 42], 6.7)

    page_break(doc)
    section_title(doc, "References and review", "Source Notes and Self-Check", "External facts are separated from ICLO direction, product principles, hypotheses and requests.")
    sources = [
        "U.S. Department of Labor - group health plans may include dental and may be insured or self-funded: https://www.dol.gov/node/63394",
        "CMS - group coverage may be fully insured or self-funded: https://www.cms.gov/medical-bill-rights/help/guides/what-type-of-insurance",
        "Snowflake Editions / Business Critical: https://docs.snowflake.com/en/user-guide/intro-editions",
        "Snowflake row access policies: https://docs.snowflake.com/en/user-guide/security-row-intro",
        "Snowflake data protection policies: https://docs.snowflake.com/en/user-guide/data-protection-policies-snowsight",
        "Snowflake Access History: https://docs.snowflake.com/en/user-guide/access-history",
        "Cross-region sharing and replication: https://docs.snowflake.com/en/user-guide/secure-data-sharing-across-regions-platforms",
        "Snowflake Native App Framework: https://docs.snowflake.com/en/developer-guide/native-apps/native-apps-about",
        "Snowflake Data Clean Rooms: https://docs.snowflake.com/en/user-guide/cleanrooms/about",
        "Delta Dental independent company structure: https://www.deltadental.com/about-us/terms-of-use/",
        "BCBS independent local-company structure: https://www.bcbs.com/about-us",
        "Dentistry.One / Delta Dental of Washington: https://www.deltadentalwa.com/our-company/community/DentistryOne-Press-Release",
        "quip / Delta ecosystem signal: https://deltadental.getquip.com/",
    ]
    bullets(doc, sources)
    paragraph(doc, "Self-check: ICLO is not presented as teledentistry; the U.S. dental-benefit workflow is explained before brand examples; Snowflake's role goes beyond hosting; employer/TPA/carrier/employee roles are distinct; no first-year savings is promised; raw-image and structured-data boundaries are clear; Business Critical is not framed as automatic HIPAA compliance; federated learning is conditional; Delta and Blue are examples rather than single-account assumptions; named owners and outputs remain at close.")
    return doc


def build_korean():
    doc = Document()
    configure(doc, internal=True)
    cover(doc, internal=True)

    page_break(doc)
    section_title(doc, "대표·임원 검토 요약", "결론부터: B-prime으로 문제를 열고, A안으로 협업을 제안합니다", "최혜윤 CMO와 김준배 CEO의 피드백과 모순되지 않도록 A안을 본체로 유지했습니다.")
    data_table(doc, ["검토 항목", "이번 문서의 반영"], [
        ["메인 방향", "Snowflake와 함께 만들 근거 레이어가 중심입니다. 미국 시장 문제는 오프닝에서만 사용합니다."],
        ["청중", "Snowflake HLS·아키텍처·파트너 담당자가 1차, 실제 payer/TPA 계정팀이 후속 청중입니다."],
        ["제안 목적", "고객 소개가 아니라 영업 유형 합의, 기술 워크숍, 실제 치아보험 데이터 계정의 검증 경로 확보입니다."],
        ["표현 원칙", "치아보험을 분명히 말하고, 첫해 절감·규제 준수·데이터 권리를 과장하지 않습니다."],
        ["패널·행사 맥락", "CEO 직접 패널 참여와 A안 메인 방향에 맞는 설명 흐름입니다."],
    ], [40, 132], 7.9)
    paragraph(doc, "ICLO의 핵심 정의: 직원이 치아보험 혜택을 이해하고, 네트워크 안의 적절한 치과를 찾고, 필요한 지원을 요청하도록 돕습니다. 그 뒤의 이용 여정을 치아보험 청구 데이터로 확인해 기업에는 집계된 근거만 제공합니다.")
    callout(doc, "쉽게 말하면", "앞에서는 직원의 치아보험 이용을 돕고, 뒤에서는 기업이 믿고 판단할 수 있는 데이터 근거를 만듭니다.", PALE_BLUE, SKY)
    callout(doc, "확정 사실", "핀테크지원센터 Fintech Cube 9기와 Shinhan Future's Lab 12기에 이미 선정되었고, 금융사 협업 과제를 발굴·검증 중입니다. 완료된 제휴나 성과로 확대해 말하지 않습니다.", PALE_TEAL, TEAL)

    page_break(doc)
    section_title(doc, "용어 먼저 보기", "미국 치아보험과 데이터 협업의 기본 용어", "국내 보험·금융 경험이 없어도 뒤 내용을 읽을 수 있도록 핵심 단어만 먼저 설명합니다.")
    data_table(doc, ["용어", "쉬운 설명"], [
        ["Employer-sponsored dental benefits", "기업이 임직원에게 제공하는 치아보험·치과복지 혜택입니다."],
        ["Payer / carrier", "치아보험 플랜을 운영하고 보험금을 지급하는 보험사 또는 보험 플랜입니다."],
        ["TPA", "기업을 대신해 가입 자격과 청구 처리를 운영하는 제3자 관리자입니다. 보험 위험을 직접 부담하지 않을 수도 있습니다."],
        ["Self-funded / ASO", "기업이 보험 비용 위험을 부담하고, 보험사·TPA는 운영만 맡는 구조입니다."],
        ["Eligibility", "어떤 직원과 가족이 어느 기간에 치아보험 대상인지 나타내는 가입 자격입니다."],
        ["Allowed / plan-paid / employee OOP", "허용된 진료비 / 보험자가 지급한 금액 / 직원이 직접 부담한 금액입니다."],
        ["Claims run-out", "진료가 끝난 뒤에도 청구가 늦게 들어오고 수정되는 기간입니다."],
        ["BAA", "개인건강정보 처리 책임을 계약으로 정하는 미국의 사업자 간 계약입니다."],
    ], [52, 120], 7.7)
    callout(doc, "주의", "Snowflake의 기술 기능이 있어도 ICLO 전체의 HIPAA·FDA·법률 문제가 자동으로 해결되지는 않습니다.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "미국 시장 문제", "치아보험에 가입해도 실제 진료 이용은 여전히 복잡합니다", "가입 여부보다 네트워크, 플랜 조건, 실제 예약 가능성을 함께 이해해야 합니다.")
    process_row(doc, [
        ("치아보험 가입", "직원·가족이 가입 대상인가"),
        ("네트워크", "어느 치과를 이용할 때 비용이 유리한가"),
        ("플랜 조건", "공제액·본인부담률·연간 한도"),
        ("예약 가능성", "치과 목록과 실제 예약이 일치하는가"),
        ("실제 이용", "예상 본인부담금을 이해했는가"),
    ])
    bullets(doc, [
        "네트워크 밖 치과를 절대 갈 수 없다는 뜻은 아닙니다. 다만 플랜에 따라 선택이 제한되거나 직원 부담이 커질 수 있습니다.",
        "치과 목록에 있다고 해서 지금 예약할 수 있다는 뜻은 아닙니다.",
        "진료 전 예상금액과 진료 후 보험 설명서(EOB)·청구 결과가 달라질 수 있습니다.",
        "기업 계약은 데이터 협의의 출발점이 될 수 있지만, 기업의 데이터 권리와 TPA·보험사의 협조가 필요합니다.",
    ])
    callout(doc, "직원 한 명의 관점", "회사 치아보험에 가입했어도 네트워크 치과 확인, 예상 본인부담금 이해, 실제 예약, 진료 후 청구 반영까지 거쳐야 혜택을 제대로 썼는지 알 수 있습니다.", PALE_BLUE, SKY)
    paragraph(doc, "데이터도 한곳에 모여 있지 않습니다. 기업은 가입 대상과 플랜 계약을, 보험사·TPA는 네트워크와 청구를, 치과는 진료·청구 정보를, ICLO 앱은 직원의 요청과 행동 기록을 각각 보유할 수 있습니다.")

    page_break(doc)
    section_title(doc, "ICLO 제품", "ICLO는 치아보험 정보와 직원의 자발적 행동을 연결합니다", "치아 사진을 찍는 앱 하나가 아니라, 임직원 치과복지 이용과 청구 확인을 연결하는 흐름입니다.")
    data_table(doc, ["구성", "직원이 보는 것", "기업·데이터 측면"], [
        ["직원 경험", "불편 사항 입력, 혜택 이해, 선택적 구강 사진, 치과 찾기, 예약 지원 요청", "직원 행동과 동의 시각을 기록합니다."],
        ["혜택·치과 정보", "내 가입 자격, 네트워크, 플랜 조건, 가능한 예약 정보", "계약으로 얻을 수 있는 범위만 사용합니다."],
        ["청구 근거", "직원에게 복잡한 분석을 보여주는 것이 목적은 아닙니다.", "가입자 월, 예방진료 완료, 허용액·보험자 지급·본인부담, 청구 지연을 기업 단위로 집계합니다."],
        ["그림자 모드 신호", "초기 질병 확률을 직원에게 보여주지 않습니다.", "모델 신호가 특정 치과·응급도·치료경로를 자동 결정하지 않습니다."],
    ], [35, 67, 70], 7.5)
    paragraph(doc, "ICLO는 원격치과 상담 앱, 소비자용 진단 AI, 직원 위험점수·인수 도구, 치과 소개 수수료 마켓플레이스, 첫해 비용절감 보장 상품으로 제안하지 않습니다.")

    page_break(doc)
    section_title(doc, "데이터 구조", "Snowflake는 데이터와 책임을 분리해 연결하는 근거 레이어입니다", "가입·플랜·이용·청구 데이터는 함께 분석하되 원본 사진과 직원별 건강 신호는 기업 화면에서 분리합니다.")
    paragraph(doc, "구조화된 근거 데이터", "구조화된 근거 데이터")
    process_row(doc, [
        ("앱 이용 기록", "동의와 직원 행동"),
        ("인사시스템 가입 자격", "직원·가족 분모"),
        ("플랜·치과 정보", "네트워크와 혜택 조건"),
        ("TPA·보험사 청구", "치아보험 청구 항목과 지연"),
        ("Snowflake 근거 레이어", "공통 모델·품질 대조·집계 결과"),
    ], [PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_TEAL])
    paragraph(doc, "원본 구강 이미지의 별도 경로", "원본 구강 이미지의 별도 경로")
    process_row(doc, [
        ("원본 사진", "통제된 미국 저장소 / 개인건강정보 보관영역"),
        ("그림자 모드 분석", "모델 엔드포인트와 검증은 별도 책임"),
        ("Snowflake에는", "URI·메타데이터·모델 버전·승인 신호만"),
        ("기업 화면에는", "집계 결과만; 직원별 경로·건강 신호 없음"),
    ], [PALE_CORAL, PALE_AMBER, PALE_BLUE, PALE_TEAL])
    callout(doc, "Business Critical + BAA", "Snowflake 안에서 개인건강정보를 처리할 때 필요한 플랫폼 전제입니다. 이것만으로 ICLO 전체의 HIPAA 준수가 완성되는 것은 아닙니다.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "경제성", "첫해 치아보험 청구비용 절감을 약속하지 않습니다", "예방진료와 미치료 상태 발견이 늘면 첫해 보험자 지급액이 먼저 증가할 수 있습니다.")
    data_table(doc, ["구간", "일어날 수 있는 변화", "해석"], [
        ["1년차", "예방진료 이용 증가, 미치료 상태 발견, 추가 치료", "허용액·보험자 지급액·직원 본인부담금이 서로 다르게 움직일 수 있습니다."],
        ["2-3년차", "치료 구성이 바뀌는지와 늦게 들어오는 청구가 보이기 시작", "방향은 고객 데이터로 검증해야 합니다."],
        ["이직률이 낮은 기업", "같은 기업이 여러 해 동안 변화를 볼 가능성이 큼", "장기 가치 논리가 상대적으로 적합합니다."],
        ["이직률이 높은 기업", "신규 입사자를 통해 미치료 수요가 다시 유입", "단기 절감보다 직원 경험·채용·유지 가치가 더 적합할 수 있습니다."],
    ], [31, 78, 63], 7.6)
    callout(doc, "J-curve", "처음에는 비용이 오르고 이후 변화 방향을 관찰하는 곡선 가설입니다. 실제 ICLO 성과값이 아닙니다.", PALE_AMBER, AMBER)
    paragraph(doc, "Snowflake에서 직원·가족 연령, 근무지, 네트워크 겹침, 예약 가능성, 이직률, 플랜 구조, 과거 청구, 본인부담 구분, 미치료 수요 가정, 청구 지연을 입력해 1-3년차와 불확실성을 비교하려는 분석 사례입니다.")

    page_break(doc)
    section_title(doc, "책임 경계", "Snowflake가 할 수 있는 일과 할 수 없는 일을 구분합니다", "Snowflake는 데이터 통제와 협업을 지원하지만 ICLO의 법률·임상·규제 결론을 대신 내리지 않습니다.")
    data_table(doc, ["Snowflake가 지원 가능한 범위", "ICLO·고객이 책임질 범위", "Snowflake에 요구하지 않는 판단"], [
        ["구조화 데이터 저장·계산, 공통 모델, 품질 대조", "원천 데이터 연결, 제품 로직, 실제 운영", "기업·TPA의 데이터 권리"],
        ["권한, 마스킹, 행 단위 정책, 접근 기록", "동의 화면, 목적 정보, 사고 대응", "동의가 법적으로 충분한지"],
        ["기업별 집계 화면, 청구로 확인한 결과 계산", "성과 귀속 로직, 고객 설명", "첫해 절감 효과 검증"],
        ["같은 리전 공유 또는 통제된 파일·API 수집", "연결기와 원천 시스템 운영", "FDA 상태, 진료 연계 면허·법률"],
        ["규모별 Snowflake 사용량과 크레딧 이후 비용 산정", "상업 가정과 실제 사용 모니터링", "ICLO 전체의 HIPAA 준수 결론"],
    ], [58, 58, 56], 7.3)
    paragraph(doc, "Native App, Clean Room, 연합학습은 첫 기업 파일럿의 필수 조건이 아닙니다. 여러 기관이 실제로 협업하거나 반복 배포가 필요해질 때 다시 검토합니다.")

    page_break(doc)
    section_title(doc, "대시보드", "성과 숫자보다 먼저 데이터의 기준을 보여줍니다", "현재 화면은 합성 데이터 예시이며 실제 성과를 주장하는 자료가 아닙니다.")
    data_table(doc, ["화면에 보여줄 항목", "쉬운 의미"], [
        ["Synthetic data - illustrative only", "실제 고객 성과가 아니라 설명용 데이터라는 표시입니다."],
        ["Employees / members", "앱을 쓰는 직원 수와 보험 비용 계산에 쓰는 전체 가입자 수를 구분합니다."],
        ["가입 자격·청구 기준일", "어느 시점까지의 데이터인지 보여줍니다."],
        ["Claims lag / completeness", "아직 들어오지 않은 청구와 데이터 누락 가능성을 보여줍니다."],
        ["집계만, n >= 20, 개인 PHI 없음", "작은 집단을 숨기고 기업이 직원 개인의 건강정보를 보지 않게 합니다."],
    ], [62, 110], 7.8)
    callout(doc, "현장 설명", "이 화면은 성과 주장이 아니라 접근성, 데이터 품질, 프라이버시가 어떻게 보이는지를 설명하는 데모입니다.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "미국 치아보험 생태계", "먼저 누가 비용을 부담하고 누가 운영하는지 봐야 합니다", "기업이 복지 프로그램을 마련하고, 치아보험사나 TPA가 운영하며, 치과는 진료 후 청구를 제출합니다.")
    process_row(doc, [
        ("기업 복지 담당", "치아보험 도입·비용 구조·운영사 선택"),
        ("치아보험사 / TPA", "네트워크·혜택·가입 자격·청구 운영"),
        ("직원 / 치과", "직원이 진료를 이용하고 치과가 청구 제출"),
        ("청구 근거", "방문·시술·허용액·보험자 지급·본인부담 확인"),
    ])
    data_table(doc, ["운영 방식", "치아보험 비용을 부담하는 쪽", "실제 운영 주체", "ICLO에 중요한 이유"], [
        ["보험형 (fully insured)", "치아보험사", "주로 치아보험사", "플랜·네트워크·청구 데이터가 해당 보험사와 법인에 연결됩니다."],
        ["자가부담형 (self-funded / ASO)", "기업", "보험사 또는 독립 TPA", "운영사가 가입 자격과 청구를 가질 수 있지만 기업의 데이터 권리는 별도 확인해야 합니다."],
    ], [34, 38, 42, 58], 7.0)
    paragraph(doc, "Delta Dental과 Blue 플랜은 이 원칙을 설명하는 후단 예시입니다. 같은 전국 브랜드 안에서도 지역·member company별 법인, 구매, 시스템, 기존 업체가 다를 수 있으므로 정확한 법인부터 확인합니다. 특정 브랜드를 하나의 고객 또는 단일 시스템으로 가정하지 않습니다.")
    callout(doc, "계정별 검증 순서", "비용 구조 -> 운영사 -> 정확한 법인 -> 클라우드·리전 -> 치아보험 가입·청구 데이터 -> 기존 업체를 차례로 확인합니다.", PALE_BLUE, SKY)
    callout(doc, "금지 표현", "계정 근거가 채워지기 전에는 'Snowflake 고객인 TPA가 많다'고 말하지 않습니다.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "Snowflake 요청", "고객 소개보다 먼저 검증 경로를 요청합니다", "Snowflake GTM 담당자에게 산업 적합성, 내부 전문가 연결, 계정팀 검증, 파트너 준비 기준을 요청합니다.")
    data_table(doc, ["시점", "Snowflake에서 받고 싶은 결과물"], [
        ["즉시", "HLS 아키텍트·솔루션 엔지니어 지정, 미국 계정·리전·Business Critical·BAA 경로, 크레딧·기술지원 범위, 2.5K/10K/25K 규모 산정, 보안 기준 구조"],
        ["다음", "payer·carrier-ASO·TPA 계정 지도, 계정팀 한 곳의 구조·수요 검토, Direct Share와 SFTP/API 기준, 파트너 체크리스트, 공동판매 준비 단계"],
        ["조건부", "Native App, 다자간 Clean Room, 연합학습, 공동 사례, Snowflake 임직원 복지 설계 파트너 가능성"],
    ], [34, 138], 7.6)
    paragraph(doc, "Fintech Cube 9기와 Shinhan Future's Lab 12기 선정은 금융사 협업을 준비하는 신뢰 신호로만 사용합니다. 치아보험 상품 제조, 개별 인수, 자동 청구 심사 기능이 이미 완성되었다는 뜻으로 말하지 않습니다.")

    page_break(doc)
    section_title(doc, "90일 실행안", "기술 구조와 계정 수요를 함께 검증합니다", "공동판매 약속이 아니라 다음 단계에 들어갈 근거를 만드는 일정입니다.")
    process_row(doc, [
        ("0-30일 | 정렬", "영업 유형·담당자·미국 계정·리전·BAA 경로·아키텍처 워크숍"),
        ("31-60일 | 설계", "데이터 수집 방식·규모 산정·보안 구조·계정표 1차본"),
        ("61-90일 | 검증", "payer/TPA 계정팀 한 곳 검토·파트너 체크리스트·진행/보류 결정"),
        ("조건부 | 확장", "Native App·Clean Room·연합학습·공동 사례·설계 파트너"),
    ])
    paragraph(doc, "Snowflake에 치과 AI를 검증하거나 ICLO를 규제 준수 상태로 만들어 달라는 요청이 아닙니다. 반복 가능한 payer/TPA 데이터 협업 구조인지, ICLO가 무엇을 입증해야 상업적 협력이 가능한지 확인하려는 제안입니다.")

    page_break(doc)
    section_title(doc, "회의 도구", "90초 영문 오프닝과 45분 아젠다", "대표가 Snowflake HLS GTM 담당자와 미팅을 시작할 때 사용할 수 있는 구성입니다.")
    paragraph(doc, "English opening statement", "English opening statement")
    opening = (
        "ICLO is building an employee dental-benefit navigation and claims-verified evidence layer. In the front, we help employees understand their dental benefits, find appropriate in-network care and request appointment support. In the back, we connect eligibility, plan context, app events and dental claims so the employer can see aggregate, privacy-governed evidence - without seeing individual health signals.\n\n"
        "In a U.S. employer-sponsored dental benefit, the employer sponsors the program, while a dental carrier or TPA may administer eligibility, network rules and claims. Employees use care, providers submit claims, and those claims confirm what happened. The problem is not coverage alone: network rules, deductibles, annual maximums, provider availability and claims lag all shape whether a covered employee can use the benefit and understand the cost. We also do not promise first-year claims savings; preventive use and newly discovered treatment can raise plan-paid claims before longer-term change becomes measurable.\n\n"
        "We see Snowflake as the governed evidence and collaboration plane behind this workflow. We are not asking Snowflake to validate our dental AI, make ICLO compliant or introduce us broadly to customers. We are asking for agreement on the HLS sales play, a technical working session on the payer/TPA data pattern, and an account-by-account path to validate whether dental eligibility or claim-line data actually exists in Snowflake."
    )
    paragraph(doc, opening)
    data_table(doc, ["시간", "순서", "회의에서 남길 것"], [
        ["0-8분", "Business fit", "Snowflake HLS 영업 유형과 핵심 이야기"],
        ["8-16분", "Account ecosystem", "대상 계정 유형과 계정표에 필요한 근거"],
        ["16-29분", "Architecture", "데이터 공유·수집, 권한, PHI 경계, 비용 질문"],
        ["29-38분", "Partnership path", "스타트업·파트너·ISV 경로와 준비 기준"],
        ["38-45분", "Next actions", "담당자, 기한, 문서 산출물"],
    ], [20, 48, 104], 7.6)

    page_break(doc)
    section_title(doc, "액션 레지스터", "회의가 좋은 대화로만 끝나지 않게 합니다", "각 액션에 Snowflake 담당자, ICLO 담당자, 기한, 필요한 결과물을 남깁니다.")
    data_table(doc, ["Action", "Snowflake owner", "ICLO owner", "Due date", "Required output"], [
        ["HLS 영업 유형 합의", "HLS GTM specialist", "CEO / GTM 담당", "회의 + 3영업일", "주요·보조 sales play 문장"],
        ["기술 워크숍 팀 지정", "HLS Industry Architect", "제품 / 데이터 담당", "회의 + 5영업일", "아키텍트·SE·보안 담당자"],
        ["미국 계정·리전·Business Critical·BAA 경로", "계정 / 보안 담당", "보안 담당", "회의 + 10영업일", "개설 절차와 전제조건 문서"],
        ["2.5K / 10K / 25K 규모 산정", "Solution Engineer", "데이터 담당", "30일", "가정·크레딧·단위 비용"],
        ["계정 매트릭스 1차본", "HLS GTM + 계정팀", "GTM 담당", "45일", "법인·치아보험 데이터·리전·담당·준비도"],
        ["payer/TPA 계정팀 1곳 검증", "Named account owner", "CEO / 제품 담당", "75일", "구조·미충족 수요 의견"],
        ["파트너 경로 결정", "Startup / Partner lead", "CEO / 파트너십", "90일", "체크리스트와 진행 / 보류 결정"],
    ], [48, 32, 30, 28, 42], 6.5)

    page_break(doc)
    section_title(doc, "검수", "최혜윤·김준배 리뷰와의 정합성", "아래 항목은 모두 Yes로 확인했습니다.")
    data_table(doc, ["검수 질문", "판정", "반영 위치"], [
        ["A안이 메인이고 B-prime은 오프닝에만 쓰였는가?", "Yes", "시장 문제 후 Snowflake 협업·근거 레이어로 전환"],
        ["치아보험과 임직원 복지 맥락이 분명한가?", "Yes", "표지, 문제, 제품, 경제성, 계정 요청 전반"],
        ["ICLO가 또 하나의 원격치과 앱처럼 보이지 않는가?", "Yes", "직원 이용 지원 + 청구 근거 + 기업 집계로 정의"],
        ["첫해 비용절감을 약속하지 않는가?", "Yes", "J-curve와 계약 원칙"],
        ["원본 사진과 청구 데이터 경계가 분명한가?", "Yes", "별도 이미지 경로와 Snowflake 메타데이터 경로"],
        ["Snowflake가 HIPAA·FDA를 자동 완성한다고 말하지 않는가?", "Yes", "Business Critical·BAA 설명과 책임 구분"],
        ["미국 치아보험 구조를 설명한 뒤 Delta·Blue를 예시로만 쓰는가?", "Yes", "참여자·비용 구조·데이터 흐름 설명 뒤 배치"],
        ["회의 후 담당자와 산출물이 남는가?", "Yes", "90일 계획과 액션 레지스터"],
    ], [108, 18, 46], 7.2)
    paragraph(doc, "추가 확인: HomeDen PMS/EMR/CRM과 보험상품 제조·개별 인수·청구심사 모델은 이번 Snowflake 제안의 메인 범위에서 제외했습니다. 해당 내용은 보험사 설계 파트너, 데이터 권리, 인간 의사결정 경계, 관할별 규제 검토가 갖춰진 뒤 조건부 후속안으로 다루는 것이 CEO·CMO 리뷰 방향과 일치합니다.")

    page_break(doc)
    section_title(doc, "출처", "외부 사실과 내부 판단을 구분한 참고자료", "공식 문서는 제품 기능과 조직 구조를 확인하는 데만 사용했고, ICLO의 시장·제품 판단은 별도 범주로 표시했습니다.")
    sources = [
        "미국 노동부 - 치과 혜택을 포함할 수 있는 기업 건강 플랜과 보험형·자가부담형 구조: https://www.dol.gov/node/63394",
        "CMS - 직장 단체보험의 fully insured·self-funded 구분: https://www.cms.gov/medical-bill-rights/help/guides/what-type-of-insurance",
        "Snowflake Editions / Business Critical: https://docs.snowflake.com/en/user-guide/intro-editions",
        "Snowflake row access policies: https://docs.snowflake.com/en/user-guide/security-row-intro",
        "Snowflake data protection policies: https://docs.snowflake.com/en/user-guide/data-protection-policies-snowsight",
        "Snowflake Access History: https://docs.snowflake.com/en/user-guide/access-history",
        "Snowflake cross-region sharing: https://docs.snowflake.com/en/user-guide/secure-data-sharing-across-regions-platforms",
        "Snowflake Native App Framework: https://docs.snowflake.com/en/developer-guide/native-apps/native-apps-about",
        "Snowflake Data Clean Rooms: https://docs.snowflake.com/en/user-guide/cleanrooms/about",
        "Delta Dental independent-company structure: https://www.deltadental.com/about-us/terms-of-use/",
        "BCBS independent local-company structure: https://www.bcbs.com/about-us",
        "Dentistry.One at Delta Dental of Washington: https://www.deltadentalwa.com/our-company/community/DentistryOne-Press-Release",
        "quip / Delta ecosystem signal: https://deltadental.getquip.com/",
        "내부 자료: 사용자 제공 Snowflake 외부 briefing 요구사항, 2026-08-05; CEO 김준배·CMO 최혜윤 이메일 피드백, 2026-08-05; ICLO 합성 Employer Dashboard.",
    ]
    bullets(doc, sources)
    return doc


def save(doc: Document, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


if __name__ == "__main__":
    save(build_english(), OUT / "02_EN_External/ICLO-Snowflake-Joint-Validation-Report-v4-EN-External.docx")
    save(build_korean(), OUT / "01_KO_Internal/ICLO-Snowflake-Joint-Validation-Report-v4-KO-Internal.docx")
    print("reports built")
