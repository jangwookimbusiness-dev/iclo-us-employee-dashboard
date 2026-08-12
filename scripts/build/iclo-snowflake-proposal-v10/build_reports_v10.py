from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor

ROOT = Path("/Users/jk0307/Documents/GitHub/iclo/iclo-us-employee-dashboard")
OUT = ROOT / "output/proposal-v10"
ICLO_LOGO = ROOT / "output/booth/ICLO-Logo-Color-Transparent.png"
SNOW_LOGO = ROOT / "output/booth/assets/Snowflake-Corporate-Logo-Blue-Transparent-v3.png"

NAVY = "1D3155"
TEAL = "008A94"
SKY = "29B5E8"
CORAL = "FF6F73"
AMBER = "E09A19"
GRAY = "63738A"
LIGHT = "F3F7FA"
PALE_BLUE = "EAF7FC"
PALE_TEAL = "EAF6F5"
PALE_CORAL = "FFF0F0"
PALE_AMBER = "FFF7E8"
WHITE = "FFFFFF"
ACTIVE_FONT = "Arial"

OPENING_EN = (
    "ICLO is building an employee dental-benefit navigation and claims-verified evidence layer. "
    "We are already receiving Snowflake startup support, and this proposal converts that relationship into a structured 90-day joint validation.\n\n"
    "The PoC is explicitly for a U.S. self-funded employer dental benefit. In the front, employees understand coverage, find plan-matching in-network options and request support. In the back, eligibility, plan context, app events and dental claims become aggregate evidence for employer decisions without exposing employee-level health signals. The employer bears dental claim risk, while a dental carrier or TPA may administer eligibility, network rules and claims.\n\n"
    "We will not promise first-year savings: preventive use and newly discovered treatment may raise plan-paid claims before longer-term treatment-mix change is measurable. We see Snowflake as the governed evidence and collaboration plane. We are not asking Snowflake to validate our dental AI, make ICLO compliant or provide generic customer introductions. We are asking the current Korea startup sponsor to route us to the appropriate U.S. HLS owners, align the sales play, validate the architecture and test one account-team use case. The goal is to learn what ICLO must prove for a repeatable, commercially real payer/TPA collaboration pattern."
)

AGENDA_EN = [
    ["0-8", "Business fit", "Primary HLS sales play and strongest Snowflake story"],
    ["8-16", "Account ecosystem", "Account types and evidence required for a named-account map"],
    ["16-29", "Architecture", "Share/ingestion, tenancy, PHI boundary and sizing questions"],
    ["29-38", "Partnership path", "Startup/partner/ISV path and readiness criteria"],
    ["38-45", "Next actions", "Named owners, dates and required written outputs"],
]

AGENDA_KO = [
    ["0-8분", "Business fit", "주요 HLS 영업 유형과 핵심 이야기"],
    ["8-16분", "Account ecosystem", "대상 계정 유형과 계정표에 필요한 근거"],
    ["16-29분", "Architecture", "공유·수집, 권한, PHI 경계, 규모 산정 질문"],
    ["29-38분", "Partnership path", "스타트업·파트너·ISV 경로와 준비 기준"],
    ["38-45분", "Next actions", "담당자, 기한, 필요한 서면 산출물"],
]

ACTIONS_EN = [
    ["Confirm support relationship and route to U.S. HLS", "Current Korea startup sponsor", "CEO / GTM lead", "Meeting + 3 business days", "Official program name, sponsor and U.S. HLS route"],
    ["Confirm HLS sales play and working-session team", "HLS GTM / Industry Architect", "CEO / product lead", "Meeting + 5 business days", "Primary/secondary play plus architect, SE and security contacts"],
    ["Confirm U.S. account, region, Business Critical and BAA path", "Account / security owner", "Security lead", "Meeting + 10 business days", "Written decision path and prerequisites"],
    ["Confirm startup credits and technical support", "Startup / Partner lead", "Finance / operations", "Meeting + 10 business days", "Amount, term, exclusions and support scope in writing"],
    ["Run 2.5K / 10K / 25K workload sizing", "Solution Engineer", "Data lead", "Day 30", "Sizing inputs and pre-credit run-rate economics; credits shown separately as a time-limited sensitivity"],
    ["Document the dental data-rights matrix", "Not a Snowflake action", "Security / legal lead", "Day 30", "Rights by source, field, purpose, recipient, retention and onward-sharing limit"],
    ["Accept the security reference architecture", "HLS architect / security owner", "Security lead", "Day 60", "Identities, data classification, encryption and keys, logging, retention, incident response, environments, named approver"],
    ["Complete one account diligence row", "HLS GTM + account owner", "GTM lead", "Day 30", "Entity, dental data, region, incumbent and readiness"],
    ["Build named-account matrix v1", "HLS GTM + account teams", "GTM lead", "Day 45", "Account-level dental workload and validation readiness"],
    ["Run one account-team use-case review", "Named account owner", "CEO / product lead", "Day 75", "Written architecture and unmet-need feedback"],
    ["Decide partner path", "Startup / Partner lead", "CEO / partnerships", "Day 90", "Checklist and proceed/defer decision"],
]

ACTIONS_KO = [
    ["기존 지원 관계 확인 및 미국 HLS 라우팅", "현재 한국 스타트업 스폰서", "CEO / GTM 담당", "회의 + 3영업일", "공식 프로그램명·스폰서·미국 HLS 연결 경로"],
    ["HLS 영업 유형·기술 검증팀 확정", "HLS GTM / Industry Architect", "CEO / 제품 담당", "회의 + 5영업일", "주요·보조 유형과 아키텍트·SE·보안 담당"],
    ["미국 계정·리전·Business Critical·BAA 경로", "계정 / 보안 담당", "보안 담당", "회의 + 10영업일", "개설 절차와 전제조건 문서"],
    ["스타트업 크레딧·기술지원 확인", "Startup / Partner lead", "재무 / 운영 담당", "회의 + 10영업일", "금액·기간·제외사항·지원 범위 서면"],
    ["2.5K / 10K / 25K 규모 산정", "Solution Engineer", "데이터 담당", "30일", "사이징 입력값과 크레딧 반영 전 정상 단가. 크레딧은 기간 한정 민감도로 별도 표기"],
    ["치아보험 데이터 권리 매트릭스 문서화", "Snowflake 요청 항목 아님", "보안 / 법무 담당", "30일", "원천·항목·목적·수령자·보관기간·재공유 제한"],
    ["보안 기준 구조 승인", "HLS 아키텍트 / 보안 담당", "보안 담당", "60일", "계정 체계·데이터 등급·암호화와 키·로깅·보관·사고 대응·환경 구분·승인자 지정"],
    ["계정 검증 행 1건 작성", "HLS GTM + 계정 담당", "GTM 담당", "30일", "법인·치아보험 데이터·리전·기존 솔루션·준비도"],
    ["계정 매트릭스 1차본", "HLS GTM + 계정팀", "GTM 담당", "45일", "계정별 치아보험 워크로드·검증 준비도"],
    ["계정팀 사용사례 검토 1건", "지정 계정 담당", "CEO / 제품 담당", "75일", "구조·미충족 수요에 대한 서면 의견"],
    ["파트너 경로 결정", "Startup / Partner lead", "CEO / 파트너십", "90일", "체크리스트와 진행 / 보류 결정"],
]


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])


def set_run(run, size=None, bold=None, color=None, font=None):
    font = font or ACTIVE_FONT
    run.font.name = font
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        rfonts.set(qn(f"w:{key}"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def configure(doc: Document, internal: bool):
    global ACTIVE_FONT
    ACTIVE_FONT = "Noto Sans KR" if internal else "Arial"
    section = doc.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(18)
    section.bottom_margin = Mm(16)
    section.left_margin = Mm(18)
    section.right_margin = Mm(18)
    section.header_distance = Mm(7)
    section.footer_distance = Mm(7)

    normal = doc.styles["Normal"]
    normal.font.name = ACTIVE_FONT
    for key in ("ascii", "hAnsi", "eastAsia", "cs"):
        normal._element.rPr.rFonts.set(qn(f"w:{key}"), ACTIVE_FONT)
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.12

    for name, size, color in [("Title", 28, NAVY), ("Heading 1", 21, NAVY), ("Heading 2", 14, TEAL), ("Heading 3", 11.5, NAVY)]:
        style = doc.styles[name]
        style.font.name = ACTIVE_FONT
        for key in ("ascii", "hAnsi", "eastAsia", "cs"):
            style._element.rPr.rFonts.set(qn(f"w:{key}"), ACTIVE_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    header = section.header
    t = header.add_table(1, 3, width=Mm(174))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.columns[0].width = Mm(28)
    t.columns[1].width = Mm(118)
    t.columns[2].width = Mm(28)
    t.cell(0, 0).paragraphs[0].add_run().add_picture(str(ICLO_LOGO), width=Mm(24))
    p = t.cell(0, 1).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 대외판 머리글은 'REPORT'다. 'PROPOSAL'로 두면 Proposal 덱과 구분되지 않는다.
    r = p.add_run("ICLO × SNOWFLAKE  |  " + ("INTERNAL REVIEW" if internal else "JOINT VALIDATION REPORT"))
    set_run(r, 9.0, True, GRAY)
    t.cell(0, 2).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    t.cell(0, 2).paragraphs[0].add_run().add_picture(str(SNOW_LOGO), width=Mm(24))
    for cell in t.rows[0].cells:
        set_cell_margins(cell, 0, 0, 40, 0)

    footer = section.footer
    footer.paragraphs[0].text = ""
    ft = footer.add_table(1, 2, width=Mm(174))
    ft.alignment = WD_TABLE_ALIGNMENT.CENTER
    ft.autofit = False
    ft.columns[0].width = Mm(140)
    ft.columns[1].width = Mm(34)
    p = ft.cell(0, 0).paragraphs[0]
    r = p.add_run("ICLO  |  August 2026")
    set_run(r, 8.4, False, GRAY)
    p2 = ft.cell(0, 1).paragraphs[0]
    add_page_number(p2)
    for cell in ft.rows[0].cells:
        set_cell_margins(cell, 0, 0, 0, 0)


def page_break(doc: Document):
    doc.add_page_break()


def eyebrow(doc: Document, text: str, color=TEAL):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text.upper())
    set_run(r, 10.2, True, color)


def title(doc: Document, text: str, subtitle: str | None = None):
    p = doc.add_paragraph(style="Title")
    p.paragraph_format.space_after = Pt(8)
    p.add_run(text)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(12)
        r = p2.add_run(subtitle)
        set_run(r, 12.5, False, GRAY)


def section_title(doc: Document, eyebrow_text: str, heading: str, takeaway: str):
    eyebrow(doc, eyebrow_text)
    p = doc.add_paragraph(style="Heading 1")
    p.paragraph_format.space_after = Pt(5)
    p.add_run(heading)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(takeaway)
    set_run(r, 11.5, False, GRAY)


def paragraph(doc: Document, text: str, bold_lead: str | None = None):
    p = doc.add_paragraph()
    if bold_lead and text.startswith(bold_lead):
        r = p.add_run(bold_lead)
        set_run(r, 11.2, True, NAVY)
        r = p.add_run(text[len(bold_lead):])
        set_run(r, 11.2, False, NAVY)
    else:
        r = p.add_run(text)
        set_run(r, 11.2, False, NAVY)
    return p


def bullets(doc: Document, items: Iterable[str], color=NAVY):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Mm(5)
        p.paragraph_format.first_line_indent = Mm(-2.5)
        r = p.add_run(item)
        set_run(r, 10.9, False, color)


def callout(doc: Document, label: str, text: str, fill: str, accent: str):
    table = doc.add_table(1, 2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Mm(43)
    table.columns[1].width = Mm(129)
    c0, c1 = table.rows[0].cells
    shade(c0, accent)
    shade(c1, fill)
    for c in (c0, c1):
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c, 130, 150, 130, 150)
    p = c0.paragraphs[0]
    r = p.add_run(label.upper())
    set_run(r, 9.4, True, WHITE)
    p = c1.paragraphs[0]
    r = p.add_run(text)
    set_run(r, 10.4, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def process_row(doc: Document, stages: Sequence[tuple[str, str]], colors: Sequence[str] | None = None):
    colors = colors or [PALE_CORAL, PALE_BLUE, PALE_AMBER, PALE_TEAL]
    table = doc.add_table(1, len(stages))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, (head, body) in enumerate(stages):
        cell = table.cell(0, idx)
        shade(cell, colors[idx % len(colors)])
        set_cell_margins(cell, 180, 160, 180, 160)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r = p.add_run(head + "\n")
        set_run(r, 11.3, True, NAVY)
        r = p.add_run(body)
        set_run(r, 9.8, False, GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def data_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[float] | None = None, font_size=9.4):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for col, width in zip(table.columns, widths):
            col.width = Mm(width)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for idx, text in enumerate(headers):
        cell = hdr.cells[idx]
        shade(cell, TEAL)
        set_cell_margins(cell, 120, 100, 120, 100)
        r = cell.paragraphs[0].add_run(text)
        set_run(r, font_size, True, WHITE)
    for row_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            shade(cell, WHITE if row_idx % 2 == 0 else LIGHT)
            set_cell_margins(cell, 100, 100, 100, 100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            r = cell.paragraphs[0].add_run(text)
            set_run(r, font_size, False, NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def cover(doc: Document, internal: bool):
    doc.add_paragraph().paragraph_format.space_after = Pt(10)
    table = doc.add_table(1, 1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, NAVY)
    set_cell_margins(cell, 500, 450, 500, 450)
    p = cell.paragraphs[0]
    r = p.add_run("ICLO × SNOWFLAKE\n")
    set_run(r, 13.0, True, SKY)
    p.add_run("\n")
    heading = "U.S. Self-Funded Employer\nDental-Benefit PoC" if internal else "U.S. Self-Funded Employer\nDental-Benefit PoC"
    r = p.add_run(heading + "\n")
    set_run(r, 25.0, True, WHITE)
    p.add_run("\n")
    sub = "내부 검토용 | 미국 치아보험·임직원 복지·데이터 구조 설명" if internal else "Snowflake HLS GTM & Architecture Joint-Validation Proposal"
    r = p.add_run(sub + "\n")
    set_run(r, 12.0, False, WHITE)
    p.add_run("\n")
    r = p.add_run("August 2026")
    set_run(r, 10.2, False, SKY)
    doc.add_paragraph()
    callout(
        doc,
        "읽는 방법" if internal else "POSITIONING",
        "기존 Snowflake 지원 관계를 90일 공동 검증으로 전환하는 미국 self-funded employer 치아보험 PoC입니다." if internal else "This converts ICLO's existing Snowflake startup support into a structured 90-day validation.",
        PALE_BLUE,
        SKY,
    )


def build_english():
    doc = Document()
    configure(doc, internal=False)
    cover(doc, internal=False)

    page_break(doc)
    section_title(doc, "Executive decision brief", "A U.S. Self-Funded Employer Dental-Benefit PoC", "Employees use dental benefits more easily; employers gain evidence for plan, administrator, access and workforce-benefit decisions.")
    process_row(doc, [
        ("1. Business fit", "Agree on the most natural Snowflake HLS sales play"),
        ("2. Architecture", "Secure an HLS technical working session"),
        ("3. Account evidence", "Validate one employer or administrator account team before introductions"),
    ], [PALE_TEAL, PALE_BLUE, PALE_AMBER])
    paragraph(doc, "ICLO helps employees understand and use their dental benefits, navigate to plan-matching in-network options, and converts the journey into claims-verified, aggregate evidence for the employer. The current PoC scope is a U.S. self-funded employer dental benefit administered by a dental carrier or TPA.")
    data_table(doc, ["Employer decision", "Evidence needed"], [
        ["Plan + administrator fit", "Eligibility quality, benefit rules, service levels and dental claims access"],
        ["Network access", "Workforce geography, provider overlap and real appointment availability"],
        ["Utilization + cost allocation", "Preventive completion; allowed, plan-paid and employee OOP"],
        ["Evidence reliability", "Denominator, source dates, claims lag, run-out and completeness"],
        ["Workforce value", "Participation, employee experience, turnover and multi-year scenarios"],
    ], [58, 114], 9.1)
    paragraph(doc, "Short form: Employee dental-benefit navigation in the front; a governed evidence layer in the back.", "Short form:")
    callout(doc, "CONFIRMED ICLO DIRECTION", "Lead with employee dental-benefit navigation and employer aggregate evidence - not oral-image AI.", PALE_BLUE, NAVY)
    callout(doc, "PRODUCT DESIGN PRINCIPLE", "Keep model-derived signals separate from automated clinical routing; keep employee-level health signals out of employer views.", PALE_TEAL, TEAL)
    callout(doc, "HYPOTHESIS / TO VALIDATE", "The pattern can become repeatable across payer/TPA accounts that hold usable dental eligibility and claim-line data.", PALE_AMBER, AMBER)
    callout(doc, "REQUEST TO SNOWFLAKE", "Provide sales-play fit, named internal stakeholders, architecture decision support, account-team validation and written readiness criteria.", PALE_BLUE, SKY)

    page_break(doc)
    section_title(doc, "Where we actually are", "What runs today, what is designed, and what does not exist yet", "Read the rest of this document against this table. We do not describe design work as if it were running.")
    data_table(doc, ["Stage", "What", "Evidence"], [
        ["Running today", "An employer dashboard demo that computes synthetic numbers in the browser. No real data, no Snowflake connection yet.", "Public demo"],
        ["Designed", "Evidence-layer data model, identity-resolution rules, policy and data-quality contracts, employer onboarding, inference API contract", "46-page engineering design"],
        ["To build in 90 days", "The Snowflake evidence layer on real data, and a dashboard fed from it", "Execution section"],
        ["Does not exist", "Employee web app, Core AI call path, benefits gateway. Separate track, not inside the 90 days.", "Engineering design, sec. 8"],
        ["Not yet decided", "What the predictive model predicts, arm assignment, the full family and guardian flow", "Engineering design, sec. 15 and 17"],
        ["Blocking start", "HIPAA role determination, BAA chain, employer-TPA data rights, legal basis for baseline load before consent", "Start gates section"],
    ], [26, 100, 46], 9.2)
    callout(doc, "Why this page is near the front", "The thicker a design document gets, the more it reads as a system already in production. One synthetic-data demo is running. Everything else is design and plan. Blur that line and we will each expect something different in 90 days.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Market context + problem", "Why the PoC Starts with a U.S. Self-Funded Employer Dental Benefit", "Korea and the U.S. do not organize dental coverage in the same way; the target operating problem is specific to the U.S. employer-benefit lane.")
    data_table(doc, ["", "Korea: practical baseline", "U.S. target context"], [
        ["Coverage structure", "NHIS is the common public baseline; some dental services remain non-covered or out of pocket", "Adult dental is often a separate benefit with its own plan, network and cost-sharing rules"],
        ["Current PoC funding", "Not the scope of this PoC", "Self-funded employer bears claims risk; a dental carrier or TPA may administer"],
        ["Employee difficulty", "Determine covered vs. non-covered services and cost", "Understand network, deductible, coinsurance, annual maximum and actual availability"],
        ["Employer difficulty", "Not the target operating model", "Eligibility, plan, provider operations and dental claims sit across organizations"],
    ], [34, 68, 70], 8.5)
    callout(doc, "EMPLOYEE PROBLEM", "Coverage does not automatically become usable care: plan rules, network economics, provider availability and final claims can differ from expectations.", PALE_CORAL, CORAL)
    callout(doc, "EMPLOYER PROBLEM", "The employer needs evidence for plan and administrator fit, access gaps, utilization, cost allocation, claims completeness and workforce value.", PALE_TEAL, TEAL)
    callout(doc, "BOUNDARY", "The comparison is directional. It does not claim that all Korean or all U.S. dental coverage follows one model.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Market ecosystem", "How the U.S. Dental Benefit Ecosystem Fits Together", "Coverage segment, funding arrangement, administrator and beneficiary are separate questions; the current PoC covers one lane.")
    data_table(doc, ["Market / funding lane", "Who bears dental claim cost", "Typical operating parties", "Current PoC?"], [
        ["Employer fully insured", "Dental carrier / insurer", "Employer, broker, carrier, employees/dependents, dentists", "No - market context"],
        ["Employer self-funded / ASO", "Employer", "Employer, broker, dental carrier or TPA, employees/dependents, dentists", "Yes - target lane"],
        ["Individual / family", "Individual and/or insurer under the product", "Marketplace or direct issuer, member, dentist", "No - market context"],
        ["Medicaid / CHIP", "Public program under federal/state structure", "State, MCO where used, beneficiary, dentist", "No - market context"],
    ], [42, 42, 70, 26], 8.2)
    process_row(doc, [
        ("Sponsor + broker", "Select funding, benefit design and administrator"),
        ("Carrier / TPA / MCO", "Operate eligibility, benefits, network and/or claims"),
        ("Member + dentist", "Receive care and submit dental claims"),
        ("Claims + evidence", "Confirm procedures, allowed, plan-paid and OOP"),
    ])
    callout(doc, "CURRENT POC SCOPE", "U.S. self-funded employer + dental carrier/TPA + employees/dependents. Data rights, legal entity, cloud/region and claim-line availability must be validated account by account.", PALE_BLUE, SKY)

    page_break(doc)
    section_title(doc, "Product", "ICLO Connects Dental-Benefit Context to Employee-Initiated Action", "ICLO is an employee benefits workflow with a claims-verified evidence layer - not another teledentistry consultation app.")
    data_table(doc, ["Layer", "What it contains", "Operating discipline"], [
        ["Employee experience", "Self-reported concern, benefit understanding, optional oral image, provider search, requested support", "Employee action or human-assisted navigation remains explicit"],
        ["Benefit and provider context", "Eligibility, network, plan rules, deductible, coinsurance, annual maximum, available appointments", "Only where contractually obtainable"],
        ["Claims evidence", "Member-month denominator, preventive completion, allowed, plan-paid, employee OOP, treatment mix, run-out", "Employer output is aggregate only"],
        ["Shadow signal", "Image reference, capture time, model version, quality flag and signal band -- plus the person surrogate key, the capturing credential and the employer id, all of which are PHI-scoped", "No raw disease probabilities; no automated provider, urgency or treatment-pathway decision"],
    ], [34, 82, 56], 9.0)
    paragraph(doc, "ICLO is not: a direct-to-consumer diagnostic AI, an underwriting or employee-risk scoring product, a provider-referral marketplace, or a first-year claims-savings promise. HomeDen PMS/EMR/CRM, insurance-product manufacturing, individual underwriting and automated claims adjudication are outside this PoC; they are conditional follow-on work, not part of what is being validated here.", "ICLO is not:")
    callout(doc, "PROVISIONAL REDLINE", "Separate signal, navigation and referral-event tables; timestamp employee action and consent; maintain purpose-based access and model-version lineage, pending U.S. regulatory counsel review.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Architecture", "Snowflake as the Governed Evidence and Collaboration Plane", "Structured eligibility, plan, event and dental-claims data converge in Snowflake; raw oral images follow a separate controlled path.")
    paragraph(doc, "Structured evidence path", "Structured evidence path")
    process_row(doc, [
        ("App events", "Consent and employee actions"),
        ("HRIS eligibility", "Employee and dependent denominator"),
        ("Plan + provider context", "Network and benefit rules"),
        ("TPA/carrier claims", "Dental claim lines and run-out"),
        ("Snowflake evidence layer", "Canonical model, controls, aggregate outcomes"),
    ], [PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_TEAL])
    paragraph(doc, "Raw-image path", "Raw-image path")
    process_row(doc, [
        ("Raw oral image", "Controlled U.S. object storage / PHI vault"),
        ("Shadow inference", "Model endpoint and validation outside this proposal"),
        ("Snowflake", "No image binaries. Person-linked references and signals do enter"),
        ("Employer dashboard", "Aggregate views; no employee-level routing or health signal"),
    ], [PALE_CORAL, PALE_AMBER, PALE_BLUE, PALE_TEAL])
    callout(doc, "PLATFORM PREREQUISITE", "Individual-level records exist in the governed processing layer - member-months, claim lines and app events cannot be computed without them. What the controls govern is who may read them: employer-facing output is aggregate and excludes employee-level PHI. If PHI is processed in Snowflake, Business Critical and an executed BAA are platform prerequisites. They do not make ICLO as a whole HIPAA-compliant.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Evidence design", "The Dashboard Makes Data Quality and Privacy Visible", "The current dashboard is synthetic visual proof, not a performance claim.")
    data_table(doc, ["Visible control", "Why the buyer needs it"], [
        ["Synthetic data - illustrative only", "Prevents the demo from being mistaken for an ICLO outcome"],
        ["Employees lens / members lens", "Participation stands on plan-eligible employees at contracted employers (source: the eligibility feed supplied by the employer or its benefits administrator, typically X12 834) - the population, not the app users. Cost stands on covered member-months, which include enrolled dependents. The two denominators are never mixed."],
        ["Eligibility and claims dates", "Shows which plan period and source cutoff each result uses"],
        ["Claims lag and completeness", "Makes run-out and missing-data uncertainty visible. The 60 days on screen is a demo constant; real lag is a distribution that varies by TPA, practice, procedure and season. A lag curve needs 12-18 months of history, so in pilot year one we show the claims-received-through date rather than an estimated completeness figure"],
        ["Aggregate only; n >= 20 suppression", "Provisional disclosure-control rule. The demo enforces it in browser JavaScript; the real system enforces it in the Snowflake query engine. Note that a Snowflake aggregation policy counts rows, not distinct people, unless an ENTITY KEY is set -- one person with 20 claim lines would otherwise satisfy a group of 20. It is not by itself a HIPAA de-identification determination"],
    ], [56, 116], 9.6)
    callout(doc, "MESSAGE", "This is not a performance claim. It demonstrates how denominator, freshness, privacy and claims completeness become visible to the buyer.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "Start gates", "No real data is loaded until these four are closed", "These are contract and legal determinations, not engineering problems. They must be settled in writing before Day 30.")
    data_table(doc, ["Gate", "What has to be settled", "If it stays open"], [
        ["HIPAA roles", "Whether ICLO, the employer, the employer health plan, the TPA and Snowflake are each a covered entity, a business associate or a subcontractor. The BAA chain is built on top of that determination.", "There is no lawful basis to load PHI."],
        ["Data rights", "The data-rights clause in the employer-TPA agreement: by source, by field, by purpose, what we may receive and what we may use it for.", "We receive data we cannot use. Most delay lands here."],
        ["Baseline load before consent", "Baseline load precedes employee consent. At that moment nobody has consented, and in a pilot most people never will. Either load an aggregate-only baseline, or settle a separate legal basis in writing.", "We would process the PHI of non-consenting employees without a basis. This is structural, not an edge case."],
        ["Common identity key", "Whether a stable identifier exists on both the HR side and the TPA side to join a person across them.", "Claims cannot be linked to app usage, so claims-confirmed completion does not hold."],
    ], [26, 96, 50], 8.8)
    callout(doc, "The order matters", "Policy and access controls go in first, real data on top of them. Reversed, there is a window in which PHI exists without controls, and that window cannot be closed retroactively.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Execution", "A 90-Day Validation with Explicit Owners, Gates and Outputs", "Snowflake and ICLO actions are separated by phase; ICLO owns the weekly cadence and one decision log.")
    data_table(doc, ["Phase", "Snowflake action", "ICLO action", "Gate / output"], [
        ["Step 0 + Days 0-30\nRoute + align", "Route to U.S. HLS; confirm play, owners and account/BAA path", "Provide current-state pack; run workshop and cadence", "Sponsor, owners and route in writing"],
        ["Days 31-60\nDesign", "Size workload; co-draft security reference and account matrix", "Provide source map, ingestion test and sizing inputs", "Architecture and matrix v1 accepted"],
        ["Days 61-90\nValidate", "Sponsor one account-team review; issue partner checklist", "Run review; recommend go/no-go", "Use case and evidence minimum agreed"],
        ["Conditional\nExpand", "Assess Native App, Clean Room or federated learning only if needed", "Build only after multi-party demand is confirmed", "Trigger: account review confirms multi-party need"],
    ], [34, 52, 52, 42], 8.2)
    callout(doc, "OPERATING + CONSUMPTION MODEL", "ICLO runs weekly written status and one decision log. ICLO supplies the sizing inputs - employers, member-months, ingestion/transform/outcome-refresh frequency, data volume, warehouse size and runtime - and the Snowflake SE calculates annual credits in the Day-30 sizing session. Day 0 is written sponsor acceptance; 'Meeting + N business days' and 'Day N' count from it.", PALE_BLUE, SKY)
    paragraph(doc, "We are not asking Snowflake to validate our dental AI or make ICLO compliant. We are asking Snowflake to help determine whether this can become a repeatable, governed payer/TPA data-collaboration pattern - and what ICLO must prove to make the partnership commercially real.")

    page_break(doc)
    section_title(doc, "Economics", "The Economics Start with Higher Year-One Claims - Not a Savings Promise", "Preventive utilization and newly discovered treatment can increase dental plan-paid claims before treatment-mix change becomes measurable. The chart shows a cost path, not a return curve.")
    data_table(doc, ["Period / segment", "What may happen", "How to interpret it"], [
        ["Year 1", "Preventive visits rise; untreated conditions are discovered", "Allowed, plan-paid and employee OOP may move differently"],
        ["Years 2-3", "Treatment mix and claims run-out become more observable", "Direction remains a customer-specific hypothesis"],
        ["Low turnover", "More time for the employer to observe multi-year effects", "A multi-year value case may be more relevant"],
        ["High turnover", "Untreated need may re-enter with new hires", "Employee experience, recruitment and retention may be the stronger case"],
    ], [32, 76, 64], 9.2)
    paragraph(doc, "ICLO will not contract on a year-one savings guarantee.")
    paragraph(doc, "Planning inputs include population age, geography, network overlap, provider availability, turnover, plan design, historical claims, paid/OOP allocation, untreated-need assumptions and claims run-out. Outputs include Year 1-3 scenarios, turnover sensitivity, network-access sensitivity and explicit uncertainty.")
    callout(doc, "HYPOTHESIS", "Simulation values are planning heuristics to be calibrated with customer data - not actual ICLO outcomes.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "Responsibilities", "What Snowflake Can Support - and What Remains Outside Its Scope", "The platform can enable governed evidence and collaboration; it does not decide ICLO's intended use, data rights or legal sufficiency.")
    data_table(doc, ["Snowflake can support", "ICLO / customer responsibility", "Outside the Snowflake ask"], [
        ["Governed storage and compute; canonical model; reconciliation", "Source mapping, product logic and operational controls", "Employer/TPA data-rights determination"],
        ["Role-based access, masking, row policies and access history", "Consent UX, purpose metadata and incident operations", "Legal sufficiency of employee consent"],
        ["Aggregate outcome views and claims-confirmed calculations", "Program attribution logic and customer reporting", "First-year savings validation"],
        ["Same-region share where available; controlled staging otherwise", "Connector and source-system implementation", "Cross-region legal conclusions"],
        ["Workload sizing; pre-credit run-rate economics", "Commercial assumptions and usage monitoring", "FDA status, licensure or referral-law analysis; ICLO-wide HIPAA compliance conclusion"],
    ], [58, 58, 56], 9.0)
    paragraph(doc, "Native App, multi-party Clean Room and federated learning are conditional options. They are not prerequisites for the first employer pilot.")

    page_break(doc)
    section_title(doc, "Partnership request", "Convert Existing Snowflake Support into a Routed Validation", "Step 0 is routing: the current Korea startup sponsor names the U.S. HLS GTM and architecture owners before technical work begins.")
    data_table(doc, ["Priority", "Requested output"], [
        ["Step 0 / immediate", "Official support-program name and current sponsor; routed U.S. HLS GTM and architecture owners; U.S. account/region/Business Critical/BAA path; written credits and support scope"],
        ["Next", "Self-funded employer dental and administrator/TPA account map; one account-team use-case validation; Direct Share versus SFTP/API decision pattern; partner checklist; co-sell milestones"],
        ["Later / conditional", "Native App feasibility; multi-party Clean Room; federated-learning feasibility; joint case study; optional Snowflake employee-benefits design partnership, subject to fit"],
    ], [38, 134], 9.4)
    bullets(doc, [
        "Which Snowflake HLS sales play fits most naturally?",
        "Is the stronger Snowflake story the employee application or the governed payer/TPA evidence layer behind it?",
        "What must ICLO prove before an account team would actively introduce or co-sell?",
        "Should the path be Snowflake for Startups, Partner Network, Powered by Snowflake or an ISV/Native App route?",
    ])

    page_break(doc)
    section_title(doc, "About ICLO", "Confirmed Facts and Fields to Complete Before Release", "ICLO has defined the product/data boundary and target PoC lane; company and Snowflake account fields must be completed before external release.")
    data_table(doc, ["Field", "Current statement"], [
        ["Company", "ICLO Co., Ltd. \u00b7 Jeju, Republic of Korea \u00b7 Commercial stage in the Korean market (B2C and B2B) \u00b7 12 core team members"],
        ["Product", "HomeDen SaaS (PMS / EMR / CRM) plus employer dental-benefit navigation and claims-evidence design"],
        ["Program selections", "Selected: Fintech Cube Cohort 9 and Shinhan Future's Lab Cohort 12"],
        ["Snowflake relationship", "Existing Snowflake startup support"],
        ["Current Snowflake state", "[Account, cloud, region, credits, usage and technical-support owner: confirm]"],
        ["Next proof", "One U.S. self-funded employer dental-data path and one account-team use-case review"],
    ], [48, 124], 8.9)
    callout(doc, "RELEASE GATE", "The remaining bracketed fields are the ask in this meeting, not gaps in ICLO's own record. Korean commercial stage and program selection do not prove insurer integration, underwriting, adjudication or U.S. outcomes.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "References and review", "Source Notes and Self-Check", "External facts are separated from ICLO direction, product principles, hypotheses and requests.")
    sources = [
        "U.S. Department of Labor, EBSA - Understanding Your Fiduciary Responsibilities Under a Group Health Plan: an ERISA-covered group health plan may cover dental, and may fund benefits through a trust, by purchasing insurance, or by self-funding from the employer's general assets. https://www.dol.gov/sites/dolgov/files/EBSA/about-ebsa/our-activities/resource-center/publications/group-health-plan-fiduciary-responsibilities.pdf (section page: https://www.dol.gov/agencies/ebsa/employers-and-advisers/plan-administration-and-compliance/health-plans)",
        "U.S. Department of Labor, EBSA - Annual Report to Congress on Self-Insured Group Health Plans (2025 edition, the current one): self-insured sponsor pays claims directly; fully insured sponsor pays premiums and the insurer assumes risk. https://www.dol.gov/sites/dolgov/files/EBSA/researchers/statistics/retirement-bulletins/annual-report-on-self-insured-group-health-plans-2025.pdf (all editions: https://www.dol.gov/agencies/ebsa/about-ebsa/our-activities/resource-center/reports)",
        "HealthCare.gov - adult dental is not an essential health benefit and may be embedded or separate: https://www.healthcare.gov/coverage/dental-coverage/",
        "Medicaid.gov, preventive services - pediatric/CHIP dental requirements and state-option adult dental: https://www.medicaid.gov/medicaid/benefits/prevention",
        "American Dental Association - dental benefit plan designs: https://www.ada.org/resources/practice/dental-insurance/benefit-plan-designs",
        "Korea NHIS - national health insurance benefit framework: https://www.nhis.or.kr/english/wbheaa02600m01.do",
        "Snowflake Editions / Business Critical: https://docs.snowflake.com/en/user-guide/intro-editions",
        "Snowflake row access policies: https://docs.snowflake.com/en/user-guide/security-row-intro",
        "Snowflake data protection policies: https://docs.snowflake.com/en/user-guide/data-protection-policies-snowsight",
        "Snowflake Access History: https://docs.snowflake.com/en/user-guide/access-history",
        "Cross-region sharing and replication: https://docs.snowflake.com/en/user-guide/secure-data-sharing-across-regions-platforms",
        "Snowflake Native App Framework: https://docs.snowflake.com/en/developer-guide/native-apps/native-apps-about",
        "Snowflake Data Clean Rooms: https://docs.snowflake.com/en/user-guide/cleanrooms/about",
    ]
    bullets(doc, sources)
    paragraph(doc, "Self-check: the U.S. self-funded employer dental PoC is explicit from the opening; ICLO is not presented as teledentistry; coverage, funding and administration are separated; Snowflake's role goes beyond hosting; employer/TPA/carrier/employee roles are distinct; no first-year savings is promised; raw-image and structured-data boundaries are clear; Business Critical is not framed as automatic HIPAA compliance; federated learning is conditional; named owners and outputs remain at close.")
    return doc


def build_korean():
    doc = Document()
    configure(doc, internal=True)
    cover(doc, internal=True)

    page_break(doc)
    section_title(doc, "무엇을 제안하는가", "미국 self-funded employer 치아보험의 근거 레이어를 Snowflake 위에 함께 만듭니다", "이 문서가 요청하는 것은 고객 소개가 아니라 검증 경로입니다 — 영업 유형 합의, 기술 워크숍, 실제 치아보험 데이터 계정 한 건.")
    paragraph(doc, "임직원이 치아보험 혜택을 이해하고, 네트워크 안의 적절한 치과를 찾고, 필요한 지원을 요청하도록 돕습니다. 그 뒤의 이용 여정을 치아보험 청구 데이터로 확인해 기업에는 집계된 근거만 제공합니다.", "ICLO의 핵심 정의")
    data_table(doc, ["항목", "내용"], [
        ["PoC 대상", "미국 self-funded employer + 치아보험사·TPA. 1차 구현 대상은 임직원 본인이며, 성인 피부양자는 독립 인증·동의가 확정된 경우에만 포함합니다. 미성년·복수 보호자는 후속 범위입니다."],
        ["요청하는 것", "영업 유형 합의, 기술 워크숍, 실제 치아보험 데이터 계정의 검증 경로 확보."],
        ["말하지 않는 것", "첫해 비용 절감, 규제 준수 완료, 확보된 데이터 권리. 셋 다 아직 입증되지 않았습니다."],
    ], [30, 142], 9.5)
    callout(doc, "임직원 가치", "임직원은 치아보험을 더 쉽게 이해하고 이용할 수 있습니다.", PALE_BLUE, SKY)
    callout(doc, "기업 의사결정", "기업은 플랜·운영사 적합성, 네트워크 접근성, 예방진료 이용, 허용액·플랜 지급액·본인부담, 청구 지연·완결성, 이직률과 장기 복지 가치를 데이터로 판단합니다.", PALE_TEAL, TEAL)

    page_break(doc)
    section_title(doc, "현재 상태", "무엇이 돌아가고 무엇이 아직 없는지 먼저 밝힙니다", "이 문서의 나머지는 이 구분 위에서 읽어야 합니다. 설계와 운영을 섞어 말하지 않습니다.")
    data_table(doc, ["구분", "항목", "근거"], [
        ["지금 돌아가는 것", "임직원 대시보드 데모 — 합성 숫자를 브라우저 안에서 계산합니다. 실데이터도 Snowflake 연결도 아직 없습니다.", "공개 데모 화면"],
        ["설계를 마친 것", "근거 레이어 데이터 모델, 신원 해석 규칙, 정책·품질 계약, 기업 온보딩 절차, 추론 API 계약", "기술 설계 문서 46쪽"],
        ["90일에 지을 것", "Snowflake 위의 실데이터 근거 레이어와 실데이터 대시보드", "이 문서 90일 실행안"],
        ["아직 없는 것", "임직원 웹앱, Core AI 호출 경로, 급여 게이트웨이 — 별도 트랙이며 90일 계획에 포함되지 않습니다.", "기술 설계 문서 8절"],
        ["아직 정하지 못한 것", "예측 모델의 예측 대상, 실험군 배정, 가족·보호자 권한의 전체 흐름", "기술 설계 문서 15·17절"],
        ["착수를 막는 것", "HIPAA 역할 규정, BAA 체인, 기업–TPA 데이터 권리, 동의 이전 기준선 적재의 법적 근거", "이 문서 착수 게이트"],
    ], [30, 96, 46], 9.2)
    callout(doc, "왜 이 페이지를 앞에 두는가", "설계 문서가 두꺼워질수록 읽는 쪽은 그것이 이미 돌아가는 시스템이라고 읽습니다. 지금 돌아가는 것은 합성 데이터 데모 하나이고, 나머지는 설계와 계획입니다. 이 구분을 흐리면 90일 뒤에 서로 다른 것을 기대하게 됩니다.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "시장 특성과 문제", "한국과 미국의 치아보험 운영 구조는 다릅니다", "이번 PoC는 미국 전체가 아니라 self-funded employer의 치아보험·임직원 복지 구조를 대상으로 합니다.")
    data_table(doc, ["구분", "한국: 이해를 위한 기준", "미국 PoC 대상 구조"], [
        ["보장 기반", "국민건강보험이 공통 기반이며, 일부 치과서비스는 비급여·본인부담", "성인 치아보험이 별도 복지 플랜으로 운영되는 경우가 많음"],
        ["비용 부담", "이번 PoC의 대상 아님", "self-funded employer가 치아보험 청구 위험을 부담"],
        ["운영", "국내 의료·보험 운영 구조", "치아보험사 또는 TPA가 가입 자격·혜택·네트워크·청구를 운영할 수 있음"],
        ["임직원 문제", "급여·비급여와 예상비용 이해", "네트워크, 공제액·본인부담률·연간 한도, 실제 예약을 함께 이해해야 함"],
        ["기업 문제", "이번 PoC의 대상 아님", "가입·플랜·치과 운영·청구 데이터가 나뉘어 플랜·운영사·복지 가치 판단이 어려움"],
    ], [33, 65, 74], 8.3)
    callout(doc, "임직원 문제", "치아보험이 있어도 어느 치과를 이용할지, 얼마나 부담할지, 실제 예약이 가능한지, 청구가 어떻게 끝났는지 알기 어렵습니다.", PALE_CORAL, CORAL)
    callout(doc, "기업 문제", "기업은 플랜·운영사 적합성, 네트워크 공백, 이용률, 청구비용과 본인부담, 데이터 완결성을 한 흐름으로 보기 어렵습니다.", PALE_TEAL, TEAL)
    callout(doc, "비교의 한계", "한국과 미국의 모든 치아보험이 하나의 구조라는 뜻은 아닙니다. 현재 PoC의 운영 차이를 이해하기 위한 제한적 비교입니다.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "미국 치아보험 생태계", "보장 시장·비용 부담·운영 주체를 나누어 봐야 합니다", "fully insured·self-funded는 비용 부담 방식이고, individual·Medicaid는 보장 시장입니다.")
    data_table(doc, ["시장 / 비용 구조", "치아보험 청구비용 부담", "주요 참여자", "현재 PoC"], [
        ["기업 fully insured", "치아보험사", "기업·브로커·치아보험사·임직원·치과", "시장 맥락"],
        ["기업 self-funded / ASO", "기업", "기업·브로커·치아보험사 또는 TPA·임직원·치과", "대상 경로"],
        ["개인 / 가족", "가입자 및 상품 구조에 따른 보험사", "Marketplace 또는 직접판매 보험사·가입자·치과", "시장 맥락"],
        ["Medicaid / CHIP", "연방·주정부 제도 아래 공공재원", "주정부·MCO·수혜자·치과", "시장 맥락"],
    ], [43, 43, 68, 26], 8.2)
    process_row(doc, [
        ("기업 + 브로커", "비용 구조·플랜·운영사 선택"),
        ("보험사 / TPA / MCO", "가입 자격·혜택·네트워크·청구 운영"),
        ("임직원·가족 + 치과", "진료 이용 후 치과가 청구 제출"),
        ("청구 + 근거", "시술·허용액·플랜 지급액·본인부담 확인"),
    ])
    callout(doc, "현재 PoC 범위", "미국 self-funded employer + 치아보험사/TPA + 임직원·가족입니다. 법인, 데이터 권리, 클라우드·리전, 청구 항목 보유 여부는 계정별로 확인합니다.", PALE_BLUE, SKY)

    page_break(doc)
    section_title(doc, "ICLO 제품", "ICLO는 치아보험 정보와 직원 행동을 연결합니다", "치아 사진을 찍는 앱 하나가 아니라, 임직원 치과복지 이용과 청구 확인을 연결하는 흐름입니다.")
    data_table(doc, ["구성", "직원이 보는 것", "데이터 원천과 현재 상태"], [
        ["직원 경험", "불편 사항 입력, 혜택 이해, 선택적 구강 사진, 치과 찾기, 지원 요청", "앱 이벤트·동의 시각. 웹앱 미구축."],
        ["정적 급여 조건", "연간 한도, 공제액, 급여 구분별 부담률, 빈도 제한", "플랜 문서(SPD)에서 사람이 옮겨 적고 2인 검증. 어떤 피드로도 오지 않습니다."],
        ["실시간 잔여 한도", "지금 남은 한도, 공제 충족 여부, 남은 횟수", "X12 270/271 실시간 조회가 있을 때만. 청구 창고로 계산하지 않습니다 — 지연 때문에 틀린 금액이 됩니다. 게이트웨이 미구축."],
        ["네트워크·예약", "네트워크 치과인지 여부", "carrier 디렉터리 링크가 1차 권장. 디렉터리 정확도가 업계 고질 문제라 복제하지 않습니다. 예약은 별도 공급 계약이 있을 때만."],
        ["청구 근거", "직원에게 분석을 보여주는 것이 목적은 아닙니다.", "가입자-월, 예방진료 완료, 허용액·플랜 지급액·본인부담, 청구 지연을 기업 단위로 집계."],
        ["그림자 모드 신호", "질병 확률을 직원에게 보여주지 않습니다.", "Core AI는 밴드(LOW/MODERATE/PRIORITY)만 반환. API 계약은 설계됐고 호출 경로는 미구축."],
    ], [30, 62, 80], 8.6)
    callout(doc, "금액보다 구분", "270/271 연동 전에는 본인부담 예상액을 표시하지 않습니다. 본인부담은 부담률·공제 충족·잔여 한도에 전부 달려 있어서, 그것 없이 금액을 보여주면 정확해 보이는 틀린 숫자가 됩니다. 대신 급여 구분을 말합니다 — 예를 들어 스케일링은 예방 구분이고 이 플랜에서 100% 보장입니다.", PALE_BLUE, SKY)
    paragraph(doc, "ICLO는 원격치과 상담 앱, 소비자용 진단 AI, 직원 위험점수·인수 도구, 치과 소개 수수료 마켓플레이스, 첫해 비용절감 보장 상품으로 제안하지 않습니다. 모델 신호·이용 지원·연계 기록을 분리하고 임직원 행동과 동의 시각을 남기며 목적별 접근과 모델 버전 계보를 유지하는 설계 경계는 미국 규제 자문 검토를 전제로 한 잠정안입니다.")

    page_break(doc)
    section_title(doc, "데이터 구조", "Snowflake는 근거 데이터를 책임별로 분리합니다", "가입·플랜·이용·청구 데이터는 함께 분석하되 원본 사진과 직원별 건강 신호는 기업 화면에서 분리합니다.")
    paragraph(doc, "구조화된 근거 데이터", "구조화된 근거 데이터")
    process_row(doc, [
        ("앱 이용 기록", "동의와 직원 행동"),
        ("인사시스템 가입 자격", "직원·가족 분모"),
        ("플랜·치과 정보", "네트워크와 혜택 조건"),
        ("TPA·보험사 청구", "치아보험 청구 항목과 지연"),
        ("근거 레이어", "Snowflake 공통 모델·품질 대조·집계 결과"),
    ], [PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_BLUE, PALE_TEAL])
    paragraph(doc, "원본 구강 이미지의 별도 경로", "원본 구강 이미지의 별도 경로")
    process_row(doc, [
        ("원본 사진", "ICLO 미국 리전 저장소 / 벤더·키·보존·감사 미결"),
        ("그림자 모드 분석", "모델 엔드포인트와 검증은 별도 책임"),
        ("Snowflake에는", "사진 바이너리는 들어가지 않음. 사람·기업에 연결된 참조와 신호는 들어감"),
        ("기업 화면에는", "집계 결과만; 직원별 경로·건강 신호 없음"),
    ], [PALE_CORAL, PALE_AMBER, PALE_BLUE, PALE_TEAL])
    callout(doc, "경계는 '파생 신호만'이 아니라 '바이너리는 안 들어감'입니다", "Snowflake에 들어가는 것은 이미지 참조·촬영 시각·모델 버전·품질 통과 여부·신호 밴드에 더해 사람 대리키, 촬영한 계정, 기업 식별자입니다. 뒤의 셋이 이미지와 신호를 개인에게 연결하므로 PHI 인벤토리와 감사 범위에 포함됩니다. 정확한 경계는 원본 사진 바이너리가 Snowflake에 들어가지 않는다는 것입니다.", PALE_CORAL, CORAL)
    callout(doc, "Business Critical + BAA", "가입자-월·청구 항목·이용 기록은 개인 단위 레코드 없이 계산할 수 없으므로, 관리형 처리 레이어에는 개인 레벨 데이터가 존재합니다. 통제 대상은 데이터의 존재가 아니라 열람 권한이며, 기업에 제공되는 결과물은 집계뿐입니다. Snowflake 안에서 개인건강정보를 처리하려면 Business Critical과 BAA가 플랫폼 전제이고, 이것만으로 ICLO 전체의 HIPAA 준수가 완성되지는 않습니다.", PALE_AMBER, AMBER)
    paragraph(doc, "이 그림에서 가장 위험한 지점은 신원 해석입니다. 인사시스템의 사람, TPA의 가입자, 앱 사용자는 서로 다른 식별자 체계를 씁니다. 이것을 하나로 묶지 못하면 청구와 앱 이용을 연결할 수 없고, '청구로 확인된 완료'라는 주장 자체가 성립하지 않습니다. 잘못 묶으면 A의 청구가 B의 이용으로 집계되는데, 그 오류는 집계 화면에서 보이지 않습니다.", "신원 해석")

    page_break(doc)
    section_title(doc, "대시보드", "성과 숫자보다 먼저 데이터의 기준을 보여줍니다", "현재 화면은 합성 데이터 예시이며 실제 성과를 주장하는 자료가 아닙니다.")
    data_table(doc, ["화면에 보여줄 항목", "쉬운 의미"], [
        ["Synthetic data - illustrative only", "실제 고객 성과가 아니라 설명용 데이터라는 표시입니다."],
        ["Employees / members", "참여율의 분모는 계약 기업의 플랜 가입 자격자이며(원천은 기업 또는 benefits administrator가 제공하는 자격 파일, 통상 X12 834) 앱 사용자가 아닙니다. 앱 사용자는 그중 활성화된 비율입니다. 비용의 분모는 가족을 포함한 가입자-월입니다. 두 분모는 섞지 않습니다."],
        ["가입 자격·청구 기준일", "어느 시점까지의 데이터인지 보여줍니다."],
        ["Claims lag / completeness", "아직 들어오지 않은 청구와 데이터 누락 가능성을 보여줍니다. 화면의 60일은 데모 표시값이며 실제 지연은 TPA·치과·시술·시기별 분포입니다. 지연 곡선은 12~18개월 실적이 있어야 하므로 파일럿 1년차에는 완결성을 추정하지 않고 청구 수신 기준일만 표시하는 것이 정직합니다."],
        ["집계만, n >= 20 셀 억제", "잠정 공개 통제 규칙입니다. 지금 데모는 이것을 브라우저 JavaScript로 판정하고, 실제 시스템에서는 Snowflake 집계 정책이 쿼리 엔진 차원에서 강제합니다. 이것만으로 HIPAA 비식별 판정이 되는 것은 아닙니다."],
    ], [62, 110], 9.4)
    callout(doc, "n >= 20은 '20명'이어야 하고, 기본값은 '20행'입니다", "Snowflake 집계 정책은 ENTITY KEY를 지정하지 않으면 그룹 크기를 행 수로 셉니다. 청구 라인·앱 이벤트는 한 사람이 여러 행을 가지므로, 한 사람이 청구 20건만 만들어도 '20명' 조건을 통과합니다. 사람 대리키를 ENTITY KEY로 지정해야 서로 다른 사람 수로 셉니다. 데모의 JavaScript 억제는 사람 수로 판정하기 때문에 이 결함이 데모에서는 드러나지 않습니다 — 개통 전 두 기업 데이터로 반드시 시험해야 합니다.", PALE_CORAL, CORAL)
    callout(doc, "현장 설명", "이 화면은 성과 주장이 아니라 접근성, 데이터 품질, 프라이버시가 어떻게 보이는지를 설명하는 데모입니다. 지금 탭은 Overview·Signals·Funnel 셋이며, 개입군 대조 추세는 실험군 배정이 미해결이라 제거했습니다.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "착수 게이트", "이 넷이 닫히기 전에는 실데이터를 적재하지 않습니다", "기술 문제가 아니라 계약·법무 판단입니다. Day 30 이전에 서면으로 확정해야 합니다.")
    data_table(doc, ["게이트", "확정해야 하는 것", "닫히지 않으면"], [
        ["HIPAA 역할", "ICLO·기업·기업 건강플랜·TPA·Snowflake 각각이 covered entity인지 business associate인지 subcontractor인지. BAA 체인이 이 판정 위에 세워집니다.", "PHI를 적재할 근거가 없습니다."],
        ["데이터 권리", "기업–TPA 계약의 데이터 권리 조항. 원천별·필드별·목적별로 무엇을 받아 무엇에 쓸 수 있는지.", "데이터를 받고도 쓸 수 없습니다. 대부분 여기서 지연됩니다."],
        ["동의 이전 기준선 적재", "기준선 적재는 임직원 동의보다 앞섭니다. 그 시점에 동의한 사람은 없고 파일럿에서 끝내 동의하지 않을 사람이 다수입니다. 집계 기준선만 적재하거나, 별도 법적 근거를 서면으로 확정해야 합니다.", "비동의자의 개인 건강정보를 근거 없이 처리하게 됩니다. 예외 가구가 아니라 모든 비동의자에게 구조적으로 발생합니다."],
        ["공통 식별키", "인사시스템과 TPA 양쪽에 사람을 이어붙일 안정적인 식별자가 있는지.", "청구와 앱 이용을 연결할 수 없어 '청구로 확인된 완료'가 성립하지 않습니다."],
    ], [30, 92, 50], 8.8)
    callout(doc, "순서를 바꾸면 안 됩니다", "정책과 권한을 먼저 세우고 그 위에 실데이터를 올립니다. 반대로 하면 개인건강정보가 통제 없는 상태로 존재하는 구간이 생기고, 그 구간은 사후에 없앨 수 없습니다.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "90일 실행안", "Snowflake와 ICLO의 단계별 역할을 분리합니다", "ICLO가 주간 서면 현황과 단일 의사결정 기록을 운영합니다.")
    data_table(doc, ["단계", "Snowflake 역할", "ICLO 역할", "게이트 / 산출물"], [
        ["Step 0 + 0-30일\n라우팅·정렬", "미국 HLS로 연결하고 영업 유형·담당·계정/BAA 경로 확정", "현황 자료 제공, 워크숍·주간 운영", "스폰서·담당·경로 서면 확정"],
        ["31-60일\n설계", "사용량 산정, 보안 기준 구조·계정 매트릭스 공동 작성", "원천 데이터 매핑·수집 시험·사이징 입력값 제공", "아키텍처·매트릭스 1차본 승인"],
        ["61-90일\n검증", "계정팀 검토 1건 후원, 파트너 준비 체크리스트 제공", "검토 운영, 진행/보류안 제시", "사용사례·최소 근거 합의"],
        ["조건부\n확장", "필요할 때만 Native App·Clean Room·연합학습 검토", "다자간 수요 확인 후 구축", "계정 검토에서 다자간 수요 확인"],
    ], [34, 52, 52, 42], 8.0)
    callout(doc, "운영·사용량 산정", "ICLO가 주간 서면 현황과 단일 의사결정 기록을 관리합니다. 기업 수, 가입자-월, 수집·변환·결과 갱신 빈도, 데이터 규모, 웨어하우스 크기와 실행 시간을 ICLO가 입력값으로 제공하고 Snowflake SE가 Day 30 사이징 세션에서 연간 크레딧을 산정합니다. Day 0은 스폰서가 라우팅을 서면으로 수락한 날이며 '회의 + N영업일'과 'Day N'은 같은 기준일에서 셉니다.", PALE_BLUE, SKY)
    paragraph(doc, "Snowflake에 치과 AI를 검증하거나 ICLO를 규제 준수 상태로 만들어 달라는 요청이 아닙니다. 반복 가능한 payer/TPA 데이터 협업 구조인지, ICLO가 무엇을 입증해야 상업적 협력이 가능한지 확인하려는 제안입니다.")

    page_break(doc)
    section_title(doc, "경제성", "첫해 치아보험 청구비용 절감을 약속하지 않습니다", "예방진료와 미치료 상태 발견이 늘면 첫해 플랜 지급액이 먼저 증가할 수 있습니다.")
    data_table(doc, ["구간", "일어날 수 있는 변화", "해석"], [
        ["1년차", "예방진료 이용 증가, 미치료 상태 발견, 추가 치료", "허용액·플랜 지급액·직원 본인부담금이 서로 다르게 움직일 수 있습니다."],
        ["2-3년차", "치료 구성이 바뀌는지와 늦게 들어오는 청구가 보이기 시작", "방향은 고객 데이터로 검증해야 합니다."],
        ["이직률이 낮은 기업", "같은 기업이 여러 해 동안 변화를 볼 가능성이 큼", "장기 가치 논리가 상대적으로 적합합니다."],
        ["이직률이 높은 기업", "신규 입사자를 통해 미치료 수요가 다시 유입", "단기 절감보다 직원 경험·채용·유지 가치가 더 적합할 수 있습니다."],
    ], [31, 78, 63], 9.1)
    callout(doc, "1년차 비용 곡선", "비용이 먼저 오르고 이후 방향은 고객 데이터로 확인하는 가설입니다. 축이 비용이므로 투자 회수 곡선(J-curve)과는 위아래가 반대입니다. 실제 ICLO 성과값이 아닙니다.", PALE_AMBER, AMBER)
    paragraph(doc, "Snowflake에서 직원·가족 연령, 근무지, 네트워크 겹침, 예약 가능성, 이직률, 플랜 구조, 과거 청구, 본인부담 구분, 미치료 수요 가정, 청구 지연을 입력해 1-3년차와 불확실성을 비교하려는 분석 사례입니다.")

    page_break(doc)
    section_title(doc, "책임 경계", "Snowflake의 지원 범위와 책임 경계를 구분합니다", "Snowflake는 데이터 통제와 협업을 지원하지만 ICLO의 법률·임상·규제 결론을 대신 내리지 않습니다.")
    data_table(doc, ["Snowflake가 지원 가능한 범위", "ICLO·고객이 책임질 범위", "Snowflake에 요구하지 않는 판단"], [
        ["구조화 데이터 저장·계산, 공통 모델, 품질 대조", "원천 데이터 연결, 제품 로직, 실제 운영", "기업·TPA의 데이터 권리"],
        ["권한, 마스킹, 행 단위 정책, 접근 기록", "동의 화면, 목적 정보, 사고 대응", "동의가 법적으로 충분한지"],
        ["기업별 집계 화면, 청구로 확인한 결과 계산", "성과 귀속 로직, 고객 설명", "첫해 절감 효과 검증"],
        ["Secure Data Sharing 또는 통제된 파일·API 수집. 같은 리전이 가장 단순하고, 리전·클라우드가 다르면 복제 비용과 데이터 이전을 함께 검토", "연결기와 원천 시스템 운영", "리전 간 데이터 이전의 법적 판단"],
        ["규모별 Snowflake 사용량과 크레딧 반영 전 정상 단가 산정", "상업 가정과 실제 사용 모니터링", "FDA 상태·면허·진료 연계 법률, ICLO 전체의 HIPAA 준수 결론"],
    ], [58, 58, 56], 8.8)
    paragraph(doc, "Native App, Clean Room, 연합학습은 첫 기업 파일럿의 필수 조건이 아닙니다. 여러 기관이 실제로 협업하거나 반복 배포가 필요해질 때 다시 검토합니다.")

    page_break(doc)
    section_title(doc, "Snowflake 요청", "기존 Snowflake 지원을 미국 HLS 검증으로 전환합니다", "Step 0에서 한국 담당이 미국 HLS GTM·아키텍처 담당자를 연결합니다.")
    data_table(doc, ["시점", "Snowflake에서 받고 싶은 결과물"], [
        ["Step 0 / 즉시", "공식 지원 프로그램명·현재 스폰서 확인, 미국 HLS GTM·아키텍처 담당 연결, 미국 계정·리전·Business Critical·BAA 경로, 크레딧·기술지원 범위"],
        ["다음", "self-funded employer 치아보험과 운영사·TPA 계정 지도, 계정팀 한 곳의 구조·수요 검토, Direct Share와 SFTP/API 기준, 파트너 체크리스트, 공동판매 준비 단계"],
        ["조건부", "Native App, 다자간 Clean Room, 연합학습, 공동 사례, Snowflake 임직원 복지 설계 파트너 가능성"],
    ], [34, 138], 9.1)
    paragraph(doc, "Fintech Cube 9기와 Shinhan Future's Lab 12기 선정은 금융사 협업을 준비하는 신뢰 신호로만 사용합니다. 치아보험 상품 제조, 개별 인수, 자동 청구 심사 기능이 이미 완성되었다는 뜻으로 말하지 않습니다.")

    page_break(doc)
    section_title(doc, "부록 A · ICLO 기본 정보", "확인된 사실과 외부 공유 전 반드시 채울 항목", "제품·데이터 경계와 PoC 대상은 정했지만 회사·Snowflake 계정 관련 대괄호 항목은 확인이 필요합니다.")
    data_table(doc, ["항목", "현재 문구"], [
        ["회사", "주식회사 아이클로(ICLO Co., Ltd.) · 제주 · 국내 시장 B2C·B2B 상용화 단계 · 핵심 인력 12명"],
        ["제품", "HomeDen SaaS(PMS·EMR·CRM)와 임직원 치아보험 이용지원·청구 근거 설계"],
        ["프로그램 선정", "핀테크지원센터 Fintech Cube 9기·신한퓨쳐스랩 12기 선정"],
        ["Snowflake 지원 관계", "기존 Snowflake 스타트업 지원"],
        ["현재 Snowflake 상태", "[계정·클라우드·리전·크레딧·사용량·기술지원 담당자 확인]"],
        ["다음 입증", "미국 self-funded employer 치아보험 데이터 경로 1건과 계정팀 사용사례 검토 1건"],
    ], [48, 124], 8.9)
    callout(doc, "확정 사실", "핀테크지원센터 Fintech Cube 9기와 Shinhan Future's Lab 12기에 선정되었습니다. 확인된 사실은 선정 그 자체이며, 완료된 제휴나 성과, 진행 중인 협업 과제로 확대해 말하지 않습니다.", PALE_TEAL, TEAL)
    callout(doc, "외부 공유 전 확인", "남은 대괄호 항목은 ICLO 내부 미비가 아니라 이번 회의의 요청 사항입니다. 국내 상용화 단계와 프로그램 선정은 보험사 연동·인수·심사 또는 미국 성과를 입증하지 않습니다.", PALE_AMBER, AMBER)

    page_break(doc)
    section_title(doc, "부록 B · 용어", "미국 치아보험과 데이터 협업의 기본 용어", "국내 보험·금융 경험이 없어도 본문을 읽을 수 있도록 핵심 단어를 모았습니다.")
    data_table(doc, ["용어", "쉬운 설명"], [
        ["Employer-sponsored dental benefits", "기업이 임직원에게 제공하는 치아보험·치과복지 혜택입니다."],
        ["Payer", "지급 책임이 있는 주체입니다. self-funded에서는 기업의 건강플랜 자체가 payer이고, 보험사가 아닙니다."],
        ["Carrier / TPA", "네트워크·심사·청구·지급 처리를 대행할 수 있는 관리자입니다. 위험을 지지 않습니다."],
        ["TPA", "기업을 대신해 가입 자격과 청구 처리를 운영하는 제3자 관리자입니다. 보험 위험을 직접 부담하지 않을 수도 있습니다."],
        ["Self-funded / ASO", "기업의 플랜이 청구 위험을 직접 보유하고, carrier·TPA는 관리·지급 처리만 맡는 구조입니다. 큰 손실에 대비해 stop-loss 보험을 따로 두기도 합니다."],
        ["Eligibility", "어떤 직원과 가족이 어느 기간에 치아보험 대상인지 나타내는 가입 자격입니다."],
        ["Allowed / plan-paid / employee OOP", "허용된 진료비 / 보험자가 지급한 금액 / 직원이 직접 부담한 금액입니다."],
        ["Claims run-out", "진료가 끝난 뒤에도 청구가 늦게 들어오고 수정되는 기간입니다."],
        ["BAA", "covered entity와 business associate 사이 또는 business associate와 그 subcontractor 사이에서, PHI의 허용된 이용·공개 범위와 보호 의무·사고 보고·반환 및 삭제를 정하는 계약입니다. 일반적인 사업자 간 개인정보 계약과 다릅니다."],
    ], [46, 126], 8.7)
    callout(doc, "주의", "Snowflake의 기술 기능이 있어도 ICLO 전체의 HIPAA·FDA·법률 문제가 자동으로 해결되지는 않습니다.", PALE_CORAL, CORAL)

    page_break(doc)
    section_title(doc, "부록 C · 출처", "외부 사실과 내부 판단을 구분한 참고자료", "공식 문서는 제품 기능과 조직 구조를 확인하는 데만 사용했고, ICLO의 시장·제품 판단은 별도 범주로 표시했습니다.")
    sources = [
        "미국 노동부 EBSA - Understanding Your Fiduciary Responsibilities Under a Group Health Plan: ERISA 적용 기업 건강플랜은 치과를 포함할 수 있고, 신탁·보험 구매·자가부담 중에서 재원 방식을 선택할 수 있습니다. https://www.dol.gov/sites/dolgov/files/EBSA/about-ebsa/our-activities/resource-center/publications/group-health-plan-fiduciary-responsibilities.pdf (섹션 페이지: https://www.dol.gov/agencies/ebsa/employers-and-advisers/plan-administration-and-compliance/health-plans)",
        "미국 노동부 EBSA - Annual Report to Congress on Self-Insured Group Health Plans(2025년판, 현행): self-insured 기업은 청구를 직접 부담하고 fully insured는 보험사가 위험을 부담합니다. https://www.dol.gov/sites/dolgov/files/EBSA/researchers/statistics/retirement-bulletins/annual-report-on-self-insured-group-health-plans-2025.pdf (전체 판: https://www.dol.gov/agencies/ebsa/about-ebsa/our-activities/resource-center/reports)",
        "HealthCare.gov - 성인 치과 혜택은 필수건강혜택이 아니며 별도 플랜일 수 있음: https://www.healthcare.gov/coverage/dental-coverage/",
        "Medicaid.gov - 아동·CHIP 치과 혜택과 주별 성인 치과 혜택: https://www.medicaid.gov/medicaid/benefits/prevention",
        "미국치과의사협회 - 치아보험 플랜 설계: https://www.ada.org/resources/practice/dental-insurance/benefit-plan-designs",
        "국민건강보험공단 - 한국 건강보험 급여 체계: https://www.nhis.or.kr/english/wbheaa02600m01.do",
        "Snowflake Editions / Business Critical: https://docs.snowflake.com/en/user-guide/intro-editions",
        "Snowflake row access policies: https://docs.snowflake.com/en/user-guide/security-row-intro",
        "Snowflake data protection policies: https://docs.snowflake.com/en/user-guide/data-protection-policies-snowsight",
        "Snowflake Access History: https://docs.snowflake.com/en/user-guide/access-history",
        "Snowflake cross-region sharing: https://docs.snowflake.com/en/user-guide/secure-data-sharing-across-regions-platforms",
        "Snowflake Native App Framework: https://docs.snowflake.com/en/developer-guide/native-apps/native-apps-about",
        "Snowflake Data Clean Rooms: https://docs.snowflake.com/en/user-guide/cleanrooms/about",
        "내부 자료: 사용자 제공 Snowflake 외부 briefing 요구사항, 2026-08-05; CEO 김준배·CMO 최혜윤 이메일 피드백, 2026-08-05; ICLO 합성 Employer Dashboard.",
    ]
    bullets(doc, sources)
    return doc


def save(doc: Document, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


if __name__ == "__main__":
    save(build_english(), OUT / "02_EN_External/ICLO-Snowflake-Joint-Validation-Report-v10-EN-External.docx")
    save(build_korean(), OUT / "01_KO_Internal/ICLO-Snowflake-Joint-Validation-Report-v10-KO-Internal.docx")
    print("reports built")

    # 빌드 끝단 정합성 검사 — 정본은 contracts/proposal-package-v11.yml
    import subprocess, sys
    repo = Path(__file__).resolve().parents[2]
    r = subprocess.run([sys.executable, str(repo / "scripts/check-package-consistency.py")],
                       cwd=repo)
    sys.exit(r.returncode)
